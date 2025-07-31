import os
import html
from datetime import datetime
from docx import Document
from openpyxl import load_workbook

from core.security import mask_phi, redact_log, sanitize_text, sanitize_filename
from core.error_handling import handle_error
from utils.file_utils import validate_file_size, get_session_temp_dir
from logger import logger

from prompts.prompt_factory import build_prompt
from core.prompts.demand_example import EXAMPLE_DEMAND, SETTLEMENT_EXAMPLE
from services.openai_client import safe_generate
from core.usage_tracker import check_quota_and_decrement
from services.dropbox_client import download_template_file

# === Polishing function ===
async def polish_demand_text(text: str) -> str:
    """
    Polishes the final demand letter: removes repetition, strengthens transitions, 
    and cuts unnecessary boilerplate.
    """
    try:
        if not text:
            return text

        prompt = f"""
You will receive a full draft of a demand letter. 
Your task is to produce a final version that is persuasive, complete, and professionally polished.

**Core Instructions:**
1. **Keep the overall structure and all section headings exactly as written** 
   (e.g., "Facts of the Occurrence", "Damages", "Settlement Demand"). 
   Do not rename or merge sections.
2. **Remove only true repetition or redundant phrasing**. Do NOT cut important facts, 
   emotional details, evidence references, or quantified damages. Retain the original language, tone, and legal fluency.
3. **Preserve emotional and narrative richness:** keep descriptions of panic attacks, 
   impact on education, family hardship, and quality-of-life losses. 
   These details are critical to persuasion.
4. Strengthen transitions between sections so the argument flows logically.
5. Use active voice, strong legal framing (duty → breach → causation → harm), 
   and persuasive tone throughout.
6. Ensure all damages, costs, and dollar amounts remain intact and clearly tied 
   to the requested settlement.
7. Avoid overly clinical or robotic language—make it compelling while remaining professional.
8. Do not shorten for the sake of word count – Trim only true redundancy. Keep the narrative full and persuasive, even if it is longer.
9. Preserve every evidence reference (e.g., police report, witness statements, security footage) – these are essential for credibility.
10. Explicitly connect injuries to damages – Ensure every physical, emotional, and financial harm is tied to the breach of duty and justified in the settlement demand.
11. Reinforce the settlement demand at the conclusion – End the letter with a strong, firm call for payment, explicitly referencing the total amount requested and consequences for failure to comply.
12. Clarify complex damages in plain English – If medical terms are used, follow them with a short plain-language explanation for impact (e.g., “bulging discs, which cause daily pain and limit mobility”).
13. Emphasize future impact – Ensure long-term medical needs, lost earning capacity, and ongoing emotional harm are clearly stated.
14. Never add or guess facts – Do not insert new injuries, events, or numbers that do not exist in the draft.
15. **Evoke empathy in the reader:** Intentionally heighten the emotional gravity of Jane's suffering and life disruption. Show the human cost of the defendant's negligence in a way that will pull at the heartstrings of the reader (without adding facts).
16. **Increase persuasive pressure:** Frame the narrative so that a defense attorney or insurance adjuster reading this letter would feel urgent pressure to settle. Make clear that a failure to resolve the case will lead to substantial exposure at trial.
17. Highlight credibility of your client: Emphasize their trustworthiness and the consistency of their account. Underscore how third‑party witnesses, medical records, and evidence corroborate her story to leave no doubt about liability.
18. Balance emotion with professionalism: While evoking empathy, ensure the tone is never inflammatory or overly casual. The letter should feel heartfelt yet authoritative, as if it were being presented directly to a judge or jury.
19. Underscore trial risk and jury appeal: Clearly suggest that a jury would be deeply sympathetic to the client's circumstances and likely award damages exceeding the current settlement demand. This frames settlement as the defendant’s best option.
20. Maintain the client’s dignity: When describing injuries, hardships, or suffering, avoid language that could be perceived as exploitative. The goal is to show the client's resilience while demonstrating the magnitude of what he or she has endured.

**Goal:** The final letter must read like it was personally reviewed and signed off by a senior trial attorney. 
It should be compelling, emotionally resonant, and leave the defense scrambling to settle.
    
Here is the draft demand letter to polish:
{text}
"""

        polished = await safe_generate(prompt)
        return polished.strip() if polished else text

    except Exception as e:
        logger.warning(f"[DEMAND_POLISH] Failed to polish demand text: {e}")
        return text


# === Async prompt generators (A+++ constraints applied) ===
async def generate_brief_synopsis(summary: str, full_name: str, example_text: str = None) -> str:
    try:
        if not summary:
            raise ValueError("No summary text provided for brief synopsis.")
        check_quota_and_decrement("internal", "openai_tokens", 1)

        prompt = build_prompt(
            "demand",
            "Brief Synopsis",
            summary,
            client_name=full_name,
            example=example_text,
        )
        result = await safe_generate(prompt)
        return result.strip() if result else "[Brief synopsis unavailable.]"
    except Exception as e:
        return handle_error(e, code="DEMAND_SYNOPSIS_001",
                            user_message="Failed to generate brief synopsis.", raise_it=True)


async def generate_combined_facts(summary: str, first_name: str, example_text: str = None) -> str:
    try:
        if not summary:
            raise ValueError("No summary text provided for facts section.")
        check_quota_and_decrement("internal", "openai_tokens", 1)

        prompt = build_prompt(
            "demand",
            "Facts/Liability",
            summary,
            client_name=first_name,
            example=example_text or EXAMPLE_DEMAND,
            extra_instructions="Do NOT mention damages or make any demand here. Facts only."
        )
        return await safe_generate(prompt)
    except Exception as e:
        return handle_error(e, code="DEMAND_FACTS_001",
                            user_message="Failed to generate facts section.", raise_it=True)


async def generate_combined_damages(damages_text: str, first_name: str, example_text: str = None) -> str:
    try:
        if not damages_text:
            raise ValueError("No damages text provided for damages section.")
        check_quota_and_decrement("internal", "openai_tokens", 1)

        prompt = build_prompt(
            "demand",
            "Damages",
            damages_text,
            client_name=first_name,
            example=example_text,
            extra_instructions="Do NOT re-argue liability. Summarize categories of harm, not detailed injuries."
        )
        return await safe_generate(prompt)
    except Exception as e:
        return handle_error(e, code="DEMAND_DAMAGES_001",
                            user_message="Failed to generate damages section.", raise_it=True)


async def generate_settlement_demand(summary: str, damages: str, first_name: str, example_text: str = None) -> str:
    try:
        if not summary and not damages:
            raise ValueError("Summary and damages text are missing for settlement demand.")
        check_quota_and_decrement("internal", "openai_tokens", 1)

        prompt = build_prompt(
            "demand",
            "Settlement Demand",
            f"{summary}\n\n{damages}",
            client_name=first_name,
            example=example_text or SETTLEMENT_EXAMPLE,
            extra_instructions="Do NOT repeat detailed facts or injuries. Only quantify damages and make the demand."
        )
        return await safe_generate(prompt)
    except Exception as e:
        return handle_error(e, code="DEMAND_SETTLEMENT_001",
                            user_message="Failed to generate settlement demand.", raise_it=True)


# === Template filling ===
def replace_placeholders(doc: Document, replacements: dict):
    """
    Replace inline placeholders in a Word document.
    """
    for paragraph in doc.paragraphs:
        text = paragraph.text
        for key, val in replacements.items():
            if key in text:
                text = text.replace(key, val)
        if paragraph.runs:
            paragraph.clear()
            paragraph.add_run(text)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    text = paragraph.text
                    for key, val in replacements.items():
                        if key in text:
                            text = text.replace(key, val)
                    if paragraph.runs:
                        paragraph.clear()
                        paragraph.add_run(text)


async def fill_template(data: dict, template_path: str, output_dir: str) -> dict:
    """
    Fill the demand template and return dict with paths for both unpolished and polished versions.
    """
    try:
        if not data or not isinstance(data, dict):
            raise ValueError("Input data for template filling is missing or invalid.")

        if not os.path.exists(template_path):
            template_path = download_template_file("demand", template_path, "templates_cache")

        validate_file_size(template_path)
        doc = Document(template_path)

        full_name = sanitize_text(data.get("Client Name", "")).strip()
        first_name = full_name.split()[0] if full_name else "Client"

        incident_date = data.get("IncidentDate", "")
        if isinstance(incident_date, datetime):
            incident_date = incident_date.strftime("%B %d, %Y")
        elif not incident_date:
            incident_date = "[Date not provided]"

        summary = sanitize_text(data.get("Summary", "[No summary provided.]"))
        damages = sanitize_text(data.get("Damages", "[No damages provided.]"))

        replacements = {
            "{{RecipientName}}": sanitize_text(data.get("RecipientName", "[Recipient Name]")),
            "{{ClientName}}": full_name or "[Client Name]",
            "{{IncidentDate}}": incident_date,
            "{{BriefSynopsis}}": await generate_brief_synopsis(summary, full_name, data.get("Example Text")),
            "{{Demand}}": await generate_combined_facts(summary, first_name, data.get("Example Text")),
            "{{Damages}}": await generate_combined_damages(damages, first_name, data.get("Example Text")),
            "{{SettlementDemand}}": await generate_settlement_demand(summary, damages, first_name, data.get("Example Text")),
        }

        replace_placeholders(doc, replacements)

        os.makedirs(output_dir, exist_ok=True)
        base_filename = f"Demand_{full_name}_{datetime.today().strftime('%Y-%m-%d')}"
        unpolished_path = os.path.join(output_dir, sanitize_filename(f"{base_filename}_UNPOLISHED.docx"))
        polished_path = os.path.join(output_dir, sanitize_filename(f"{base_filename}_POLISHED.docx"))

        # Save unpolished
        doc.save(unpolished_path)

        # Polish entire text and overwrite to new polished document
        full_text = "\n\n".join([p.text for p in doc.paragraphs if p.text.strip()])
        polished_text = await polish_demand_text(full_text)

        polished_doc = Document()
        for paragraph in polished_text.split("\n"):
            if paragraph.strip():
                polished_doc.add_paragraph(paragraph.strip())
        polished_doc.save(polished_path)

        logger.info(f"[DEMAND_GEN] Saved unpolished: {unpolished_path}, polished: {polished_path}")

        return {"unpolished": unpolished_path, "polished": polished_path}

    except Exception as e:
        handle_error(e, code="DEMAND_FILL_001",
                     user_message="Failed to fill demand template.", raise_it=True)


async def generate_all_demands(template_path: str, excel_path: str, output_dir: str):
    try:
        if not os.path.exists(excel_path):
            handle_error(FileNotFoundError(f"Excel file not found: {excel_path}"),
                         code="DEMAND_EXCEL_001", user_message="Excel input file not found.", raise_it=True)

        wb = load_workbook(excel_path)
        sheet = wb.active
        headers = [str(cell.value).strip() for cell in sheet[1] if cell.value]

        os.makedirs(output_dir, exist_ok=True)

        for idx, row in enumerate(sheet.iter_rows(min_row=2, values_only=True), start=2):
            data = dict(zip(headers, row))
            if not data.get("Client Name", "").strip():
                logger.warning(redact_log(mask_phi(f"[DEMAND_BATCH_SKIP] Row {idx} skipped: missing Client Name")))
                continue

            await fill_template(data, template_path, output_dir)

    except Exception as e:
        handle_error(e, code="DEMAND_BATCH_001",
                     user_message="Failed to generate demand letters from Excel.", raise_it=True)


async def generate_demand_letter(
    client_name: str,
    defendant: str,
    location: str,
    incident_date: str,
    summary: str,
    damages: str,
    template_path: str,
    output_path: str,
    example_text: str = None,
):
    try:
        data = {
            "Client Name": client_name,
            "Defendant": defendant,
            "Location": location,
            "IncidentDate": incident_date,
            "Summary": summary,
            "Damages": damages,
            "RecipientName": defendant,
            "Example Text": example_text or "",
        }
        return await fill_template(data, template_path, os.path.dirname(output_path))

    except Exception as e:
        handle_error(e, code="DEMAND_GEN_001",
                     user_message="Failed to generate demand letter.", raise_it=True)


if __name__ == "__main__":
    try:
        TEMPLATE_NAME = "demand_template.docx"
        TEMPLATE_PATH = download_template_file("demand", TEMPLATE_NAME, "templates_cache")
        EXCEL_PATH = "data_demand_requests.xlsx"
        OUTPUT_DIR = get_session_temp_dir()
        import asyncio
        asyncio.run(generate_all_demands(TEMPLATE_PATH, EXCEL_PATH, OUTPUT_DIR))
    except Exception as e:
        handle_error(e, code="DEMAND_MAIN_001",
                     user_message="Error in main demand generator run.")
