import os
import re
from datetime import datetime
from docx import Document
from openpyxl import load_workbook
from openai import OpenAI
from docx.table import _Cell
from docx.text.paragraph import Paragraph

# === Load API Key and Initialize Client ===
with open("Open AI Secret Key.txt", "r") as f:
    api_key = f.read().strip()

client = OpenAI(api_key=api_key)

# === Prompt Guidelines ===
NO_HALLUCINATION_NOTE = """
Do not fabricate or assume any facts. Use only what is provided. Avoid headings, greetings, and signoffs — the template handles those. Refer to the client by their first name only. Keep all naming, pronouns, and chronology consistent. Do not use more than one version of the incident. Do not repeat injury or treatment details across sections.
"""

LEGAL_FLUENCY_NOTE = """
Use the tone and clarity of a senior litigator. Frame facts persuasively using legal reasoning: duty, breach, causation, and harm. Eliminate redundancy, vague phrases, and casual storytelling. Frame liability clearly. Maintain formal, polished, and precise language. Quantify damages where possible. Refer to witnesses, police, and footage once. Avoid any instance of 'Jane Roe' or 'Amy' — only use the first name.
Do not restate the client’s injuries more than once. After the initial mention, refer to them only by category (e.g., “orthopedic trauma,” “soft tissue damage,” “ongoing symptoms”).

Eliminate any of the following weak or redundant phrases: “continues to uncover injuries,” “in the process of obtaining,” “we believe,” “potential footage,” or “may have been.”

Use strong, legally assertive alternatives:
- “Reports symptoms consistent with...”
- “Surveillance footage is being secured...”
- “Liability is well-supported by the available evidence...”

In the closing paragraph, avoid overexplaining. End firmly with one or two sentences:
“We invite resolution of this matter without the need for litigation. Should you fail to respond by [date], we are prepared to proceed accordingly.”

All content must sound like it was drafted for final review by a managing partner or trial attorney. Every sentence should advance legal theory, factual support, or damage justification — never simply restate.

Avoid summarizing facts multiple times. Focus instead on drawing conclusions from the established facts.
"""

BAN_PHRASING_NOTE = """
Ban any phrasing that introduces speculation or weakens factual strength. Do not use: “may,” “might,” “potential,” “appears to,” “possibly,” or “believes that.” Replace all with direct phrasing: “Jane is,” “The evidence will show,” “The footage depicts...”
"""

FORBIDDEN_PHRASES = """
Forbidden: “continues to discover injuries,” “a host of,” “significant emotional hardship,” “cannot be overlooked,” “it is clear that,” “ongoing discomfort,” “found herself,” “left her with,” “had to,” “was forced to,” “Jane was returning,” “she elected to,” “engrossed in conversation,” “was caught off guard”
"""

NO_PASSIVE_LANGUAGE_NOTE = """
Every sentence must use active voice. Eliminate all passive constructions. Do not say “was struck” or “has been advised.” Instead: “The snowplow struck Jane,” or “Jane is gathering...” 
"""

SETTLEMENT_EXAMPLE = '''
This is not a complicated case, and a jury will easily be able to conclude that the snowplow operator acted negligently and, as a result, Jane was seriously injured. Jane, like many pedestrians, had no safe option but to walk in the street. Despite this, the operator failed to see her, ultimately running her over and leaving her pinned beneath the vehicle. 

A lay jury will empathize with Jane. Many people have experienced close calls with large vehicles or have struggled to be seen in winter conditions. Moreover, the police arrived at the scene, and potential video footage may corroborate Jane’s account. The damages here are significant and will be difficult for your client to contest.

Based on similar verdicts and settlements for cases involving traumatic injuries, we are authorized to demand the sum of $100,000.00 (One Hundred Thousand Dollars) to resolve this matter. If we do not receive a timely response, we are prepared to proceed with litigation.
'''

# === OpenAI Request Handler ===
def generate_with_openai(prompt):
    try:
        response = client.chat.completions.create(
            model="gpt-4",
            messages=[
                {"role": "system", "content": "You are a senior litigation attorney drafting demand letters with formal, confident, and legally persuasive tone. Use active voice, frame liability using duty/breach/causation/harm, and eliminate repetition, vague language, or passive phrasing."},
                {"role": "user", "content": prompt}
            ],
            max_tokens=1400,
            temperature=0.4
        )
        return response.choices[0].message.content.strip()
    except Exception as e:
        print(f"OpenAI API Error: {e}")
        return "[Error generating content]"

# === Section Generators ===
def generate_brief_synopsis(summary, full_name):
    prompt = f"""
{NO_HALLUCINATION_NOTE}
{LEGAL_FLUENCY_NOTE}

Write a one-sentence introduction for a legal demand letter. Refer to the client only as {full_name}. Begin with a sentence like: "Our office represents {full_name}, who suffered..." and describe the injuries factually and concisely. Do not repeat the client's name. Do not use more than one sentence.

Summary of Incident:
{summary}
"""
    response = generate_with_openai(prompt)
    first_sentence = response.split(".")[0].strip() + "."
    return first_sentence

EXAMPLE_DEMAND = """
The snowplow operator owed a duty of care to nearby pedestrians and failed to exercise basic precautions while maneuvering in a populated area. This breach of duty was the direct and proximate cause of Jane’s orthopedic trauma and emotional harm. Accordingly, liability is clear.
"""

def generate_demand(summary, first_name):
    prompt = f"""
{NO_HALLUCINATION_NOTE}
{LEGAL_FLUENCY_NOTE}

Use this as a style guide only — DO NOT use or reference the facts or content below:
Frame the conduct in terms of:
- Duty
- Standard of Care
- Breach
- Causation
- Harm

Use the following example only for tone and structure:
{EXAMPLE_DEMAND}

Do not re-describe facts. Simply reference them. Focus on legal logic. Refer to the client as {first_name} only.

Do not re-describe injuries in detail. Reference them succinctly as “the injuries described above” or by category (e.g., orthopedic trauma, disfigurement, neurological symptoms).

Strengthen the argument.

Phrases like “may be in violation…” sound too tentative. Be more assertive: “This conduct violated the driver’s duty of care and applicable safety protocols.”

Avoid passive constructions or phrases like “Jane failed to hear...” or “there may be evidence.” Instead, use direct, confident phrasing such as:
- “Jane continues to report symptoms consistent with...”
- “We anticipate corroboration through eyewitness testimony and available security footage.”

Avoid phrasing that implies fault on the part of the client. Emphasize environmental conditions, operator inattention, or institutional duty where appropriate. If applicable, cite plausible standards or duties that may apply to snowplow operators (e.g., municipal safety protocols, duty to yield to pedestrians, or winter operation guidelines). Do not invent law.

Avoid repeating injury descriptions across sections. Reference the nature of injuries once, then analyze them legally or by category (e.g., ongoing impairment, disfigurement). Avoid vague or speculative phrases like "ongoing discovery of injuries" unless supported.

State injuries once in detail in the introduction. In subsequent sections, refer to “the injuries described above” or summarize them categorically (e.g., orthopedic trauma, disfigurement).

Use formal, clinical language. Avoid casual phrases like “a host of injuries” or informal visuals like “with only her head visible.” Prefer phrasing such as: “sustained contusions and abrasions consistent with being dragged beneath the vehicle.”

Frame elements more precisely. Emphasize that the assailant failed to do certain things that ultimately resulted in the incident. 

Refine tone to remove informality and literary phrasing.

Where possible, preview anticipated evidence more confidently (e.g., “anticipated to show…” vs. “might show…”).

Condense repetitive phrasing. 

Make the transitions and flow fluent. 

Trim redundancy.

Avoid redundancy.

Always use an active voice. 

Replace speculative language with assertive phrasing. This includes all references to anticipated or possible evidence.

Polish the legal tone throughout: elevate it to match the voice of a senior litigator rather than a paralegal summary.

Avoid passive or speculative phrases like “Jane is also discovering…” or “potentially available evidence.” Use active and confident legal language such as “Jane continues to report symptoms consistent with...” or “we anticipate corroborating this account through...” 

Where appropriate, include a nod to applicable municipal safety codes or standard snowplow operator responsibilities (e.g., failure to maintain visual awareness, reckless disregard for pedestrian safety). Do not invent laws, but use plausible phrasing like “in violation of their duty under local safety protocols.”

Avoid awkward or speculative phrases. Use polished, professional language. Prefer:
- “an impact she recalls only as...” over “an event she felt like...”
- “has reported ongoing discomfort...” over “continues to discover injuries”

This section must be written as if spoken in closing argument to a jury. Assume facts are already established and focus solely on logical causation and legal accountability. Use legal transitions such as:
- “This breach of duty was the direct and proximate cause of...”
- “Accordingly, liability is established under...” 
- “Based on these facts, recovery is warranted under...”


Summary of Incident:

Summary of Incident:
{summary}
"""
    return generate_with_openai(prompt)


def generate_damages(damages_text, first_name):
    prompt = f"""
{NO_HALLUCINATION_NOTE}
{LEGAL_FLUENCY_NOTE}

Draft a damages section for the client using only the information below:
Avoid repeating the full injury narrative. Focus on economic and non-economic harms, not injury descriptions. Where possible, suggest how the total amount may be reasonably allocated across categories like medical expenses, lost wages, and pain and suffering. Do not fabricate specifics—rely only on the content below.

Frame the impact of those injuries in terms of:
- medical disruption,
- ongoing treatment,
- financial strain, and
- emotional hardship.

Limit any phrases that describe physical injuries. Focus solely on economic hardship, quantifiable expenses, and emotional impairment. Do not reference the incident again. Focus only on aftermath.

Reference the total cost and nature of the treatment. You may group expenses by provider (e.g., ambulance, hospital, surgical care, dermatology), but do not itemize every line. Emphasize the scope and ongoing nature of treatment. Use a professional tone.

Avoid repeating injury descriptions across sections. Reference the nature of injuries once, then analyze them legally or by category (e.g., ongoing impairment, disfigurement). Avoid vague or speculative phrases like "ongoing discovery of injuries" unless supported.

State injuries once in detail in the introduction. In subsequent sections, refer to “the injuries described above” or summarize them categorically (e.g., orthopedic trauma, disfigurement).

Refine tone to remove informality and literary phrasing.

Refine tone to remove informality and literary phrasing.

Where possible, preview anticipated evidence more confidently (e.g., “anticipated to show…” vs. “might show…”).

Condense repetitive phrasing. 

Make the transitions and flow fluent. 

Trim redundancy.

Avoid redundancy.

Strengthen language around non-economic harm:
- Replace “continues to discover injuries” with “reports persistent symptoms consistent with...”
- Replace “a host of minor issues” with “ongoing physical and emotional complications related to the trauma”

Avoid passive or speculative phrases like “Jane is also discovering…” or “potentially available evidence.” Use active and confident legal language such as “Jane continues to report symptoms consistent with...” or “we anticipate corroborating this account through...” 

Use formal, clinical language. Avoid casual phrases like “a host of injuries” or informal visuals like “with only her head visible.” Prefer phrasing such as: “sustained contusions and abrasions consistent with being dragged beneath the vehicle.”

Do not reference phrases like “Substantial” and “significant” too often — vary language (e.g., “considerable,” “material,” “marked,” “lasting impact”).

Avoid awkward or speculative phrases. Use polished, professional language. Prefer:
- “an impact she recalls only as...” over “an event she felt like...”
- “has reported ongoing discomfort...” over “continues to discover injuries”

Where relevant, incorporate follow-up appointments, diagnosis details, or ongoing care plans (e.g., physical therapy, orthopedic follow-up). Justify medical costs using phrases like “based on available billing records” or “projected from current provider estimates.” If income loss is mentioned, clarify timeframe, employment type, or hourly rate if known.


{damages_text}
"""
    return generate_with_openai(prompt)

SETTLEMENT_PROMPT = f"""
{NO_HALLUCINATION_NOTE}
{LEGAL_FLUENCY_NOTE}

Use a professional, authoritative tone. Avoid casual or speculative phrasing such as “This is not a complicated case.” Prefer legally grounded language like “Liability is clear based on available evidence and anticipated discovery.” Where appropriate, reference local safety obligations or common standards of care for municipal vehicle operators.

Draft a settlement demand conclusion paragraph that persuasively states why liability is clear, what discovery would confirm, and justifies the compensation request of $100,000.
"""

def generate_settlement_demand(summary, damages, first_name):
    prompt = f"""
{NO_HALLUCINATION_NOTE}
{LEGAL_FLUENCY_NOTE}

Use a professional, authoritative tone. Avoid casual or speculative phrasing. Do not invent any facts or sources of liability. Do not reference things like "bus footage" or "witnesses" unless explicitly provided.

Draft a closing settlement paragraph for {first_name} that justifies the demand based on the facts and injuries described below. Do not add any legal theories or damages not mentioned. Reference the strength of clients position based on factual consistency, corroborating sources, and the nature of their injuries.

Refine tone to remove informality and literary phrasing.

Where possible, preview anticipated evidence more confidently (e.g., “anticipated to show…” vs. “might show…”).

Condense repetitive phrasing. 

Make the transitions and flow fluent. 

Trim redundancy.

Avoid redundancy.

Integrate a single strong sentence on why litigation risk is high if this isn’t resolved.

End with a direct call to action: “We invite resolution of this matter without the need for formal litigation. Should you fail to respond by [insert date], we are prepared to proceed accordingly.”

Incident Summary:
{summary}

Damages:
{damages}
"""
    return generate_with_openai(prompt)


# === Placeholder Replacement ===
def replace_placeholders(doc, replacements):
    def replace_in_paragraph(paragraph: Paragraph):
        full_text = paragraph.text
        for key, val in replacements.items():
            if key in full_text:
                full_text = full_text.replace(key, val)
        if paragraph.runs:
            paragraph.clear()
            paragraph.add_run(full_text)

    def replace_in_cell(cell: _Cell):
        for paragraph in cell.paragraphs:
            replace_in_paragraph(paragraph)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell)

# === Word Template Filler ===
def fill_template(data, template_path, output_path):
    doc = Document(template_path)

    full_name = data.get("Client Name", "").strip()
    first_name = "Jane"

    incident_date = data.get("IncidentDate", "")
    if isinstance(incident_date, datetime):
        incident_date = incident_date.strftime("%B %d, %Y")
    elif not incident_date:
        incident_date = "[Date not provided]"

    summary = data.get("Summary", "") or "[No summary provided.]"
    damages = data.get("Damages", "") or "[No damages provided.]"

    replacements = {
        "{{Client Name}}": full_name,
        "{{IncidentDate}}": incident_date,
        "{{Brief Synopsis}}": generate_brief_synopsis(summary, first_name),
        "{{Demand}}": generate_demand(summary, first_name),
        "{{Damages}}": generate_damages(damages, first_name),
        "{{Settlement Demand}}": generate_settlement_demand(summary, damages, first_name)
    }

    replace_placeholders(doc, replacements)

    output_filename = f"Demand_{full_name.replace(' ', '_')}_{datetime.today().strftime('%Y-%m-%d')}.docx"
    doc.save(os.path.join(output_path, output_filename))
    print(f"Generated: {output_filename}")

# === Excel Integration ===
def generate_all_demands(template_path, excel_path, output_dir):
    wb = load_workbook(excel_path)
    sheet = wb.active
    headers = [str(cell.value).strip() for cell in sheet[1]]
    os.makedirs(output_dir, exist_ok=True)

    for row in sheet.iter_rows(min_row=2, values_only=True):
        data = dict(zip(headers, row))
        fill_template(data, template_path, output_dir)

# === Main Execution ===
if __name__ == "__main__":
    TEMPLATE_PATH = "templates_demand_template.docx"
    EXCEL_PATH = "data_demand_requests.xlsx"
    OUTPUT_DIR = "output_requests"

    generate_all_demands(TEMPLATE_PATH, EXCEL_PATH, OUTPUT_DIR)
    print("All demand letters generated.")
