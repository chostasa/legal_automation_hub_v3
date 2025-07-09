from openai import APIStatusError
import time
import os
from datetime import datetime
from docx import Document
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from openai import OpenAI
import streamlit as st
import re

try:
    api_key = st.secrets["OPENAI_API_KEY"]
except Exception:
    api_key = os.getenv("OPENAI_API_KEY", "")  # fallback for local dev

client = OpenAI(api_key=api_key)

def trim_to_token_limit(text, max_tokens=12000):
    if not text: return ""
    tokens = text.split()
    if len(tokens) <= max_tokens:
        return text
    third = max_tokens // 3
    return " ".join(tokens[:third]) + "\n...\n" + " ".join(tokens[-2 * third:])

def safe_generate(fn, *args, retries=3, wait_time=10, **kwargs):
    trimmed_args = [trim_to_token_limit(arg, 6000) if isinstance(arg, str) else arg for arg in args]
    for attempt in range(retries):
        try:
            return fn(*trimmed_args)
        except APIStatusError as e:
            if e.status_code == 429 and "rate_limit_exceeded" in str(e):
                st.warning(f"Rate limit hit. Retrying in {wait_time} seconds...")
                time.sleep(wait_time)
            else:
                raise e
    raise Exception("❌ gpt-3.5-turbo rate limit error after multiple attempts.")


def generate_with_openai(prompt, model="gpt-3.5-turbo"):
    response = client.chat.completions.create(
        model=model,
        messages=[
            {"role": "system", "content": "You are a professional legal writer."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.5
    )
    return response.choices[0].message.content.strip()

def embed_quotes_in_section(text, quotes, heading="TESTIMONY"):
    """
    Append formatted quotes into a section body.
    """
    if not quotes.strip():
        return text.strip()

    return f"{text.strip()}\n\n{heading}:\n{quotes.strip()}"

def chunk_text(text, max_chars=6000):
    """
    Safely chunk long text inputs by character length to stay under GPT's context limit.
    """
    chunks = []
    start = 0
    while start < len(text):
        end = start + max_chars
        if end < len(text):
            # Avoid cutting off mid-sentence
            end = text.rfind(".", start, end) + 1 or end
        chunks.append(text[start:end].strip())
        start = end
    return [c for c in chunks if c]

# === NEW FUNCTIONS FOR PREPROCESSING TRANSCRIPTS ===

def normalize_deposition_lines(raw_text):
    """
    Attaches the current page number to each line using format 0004:12.
    """
    lines = raw_text.splitlines()
    current_page = None
    numbered_lines = []

    for line in lines:
        page_match = re.match(r"^(0\d{3})\s*$", line.strip())
        if page_match:
            current_page = page_match.group(1)
            continue

        line_match = re.match(r"^\s*(\d{1,2})\s+(.*)$", line)
        if line_match and current_page:
            line_num = int(line_match.group(1))
            content = line_match.group(2).strip()
            full_id = f"{current_page}:{line_num:02d}"
            numbered_lines.append((full_id, content))

    return numbered_lines


def merge_multiline_qas(numbered_lines):
    """
    Rebuilds full Q/A blocks including multi-line content.
    Returns a string of the cleaned transcript.
    """
    qa_blocks = []
    buffer = []
    current_type = None

    for id_line, content in numbered_lines:
        if content.startswith("Q:"):
            if buffer:
                qa_blocks.append("\n".join(buffer))
                buffer = []
            current_type = "Q"
            buffer.append(f"{id_line} {content}")
        elif content.startswith("A:"):
            if buffer:
                qa_blocks.append("\n".join(buffer))
                buffer = []
            current_type = "A"
            buffer.append(f"{id_line} {content}")
        elif current_type in {"Q", "A"}:
            buffer.append(f"{id_line} {content}")

    if buffer:
        qa_blocks.append("\n".join(buffer))

    return "\n".join(qa_blocks)



# === Prompt Guidelines ===
NO_HALLUCINATION_NOTE = """
Do not fabricate or assume any facts. Use only what is provided. Avoid headings, greetings, and signoffs — the template handles those. Refer to the client by their first name only. Keep all naming, pronouns, and chronology consistent. Do not use more than one version of the incident. Do not repeat injury or treatment details across sections.
"""

LEGAL_FLUENCY_NOTE = """
Use the tone and clarity of a senior litigator. Frame facts persuasively using legal reasoning: duty, breach, causation, and harm. Eliminate redundancy, vague phrases, and casual storytelling. Frame liability clearly. Maintain formal, polished, and precise language. Quantify damages where possible. Refer to witnesses, police, and footage once. Avoid any instance of 'Jane Roe' or 'Amy' — only use the first name.
Do not restate the client’s injuries more than once. After the initial mention, refer to them only by category (e.g., “orthopedic trauma,” “soft tissue damage,” “ongoing symptoms”).

Eliminate any of the following weak or redundant phrases: “continues to uncover injuries,” “in the process of obtaining,” “we believe,” “potential footage,” or “may have been.”

Use strong, legally assertive alternatives:
- “Reports symptoms consistent with...”
- “Surveillance footage is being secured...”
- “Liability is well-supported by the available evidence...”

In the closing paragraph, avoid overexplaining. End firmly with one or two sentences:
“We invite resolution of this matter without the need for litigation. Should you fail to respond by [date], we are prepared to proceed accordingly.”

All content must sound like it was drafted for final review by a managing partner or trial attorney. Every sentence should advance legal theory, factual support, or damage justification — never simply restate.

Avoid summarizing facts multiple times. Focus instead on drawing conclusions from the established facts.
"""

BAN_PHRASING_NOTE = """
Ban any phrasing that introduces speculation or weakens factual strength. Do not use: “may,” “might,” “potential,” “appears to,” “possibly,” or “believes that.” Replace all with direct phrasing: “Jane is,” “The evidence will show,” “The footage depicts...”
"""

FORBIDDEN_PHRASES = """
Forbidden: “continues to discover injuries,” “a host of,” “significant emotional hardship,” “cannot be overlooked,” “it is clear that,” “ongoing discomfort,” “found herself,” “left her with,” “had to,” “was forced to,” “Jane was returning,” “she elected to,” “engrossed in conversation,” “was caught off guard”
"""

NO_PASSIVE_LANGUAGE_NOTE = """
Every sentence must use active voice. Eliminate all passive constructions. Do not say “was struck” or “has been advised.” Instead: “The snowplow struck Jane,” or “Jane is gathering...”
"""

INTRO_EXAMPLE = """
{{Plaintiff}}, 42, will require a L5-S1 decompressive hemilaminectomy and microdiscectomy as a result of the Defendant driver rear ending Stan at over 50 mph on his very first trip as a trucker. Plaintiff has been given leave to pursue punitive damages against the defendants. Defendant STL Trucking has failed to produce even one witness for deposition. STL claims that no STL personnel have knowledge of the crash and that depositions are not relevant. (Ex. A, Plaintiff’s Complaint at Law)
The jury will punish defendants and STL Truckers for their conscious disregard for safety and training of the Defendant driver. Mr. Doe has permanent disc herniations at L3-4, L4-5, and L5-S1. Plaintiff will require future medical care including surgery, injections and physical therapy for the rest of his life. Mr. Doe will require $869,952.41 related to the August 8, 2018 crash.
Defendants state they have a $1,000,000.00 eroding policy for this case.  They refuse to sign an affidavit that there is excess or umbrella coverage.  Plaintiff made a pre-suit demand of $1,000,000.00.  Defendants had offered a nominal amount of $50,000.00.  Stan will begin negotiating when Plaintiff is aware of the true policy limits. Verdict potential exceeds $3,000,000.00 for a case that has life altering injuries and punitive damages.
"""


def generate_introduction(input_text, client_name):
    prompt = f"""
{NO_HALLUCINATION_NOTE}
{LEGAL_FLUENCY_NOTE}
{BAN_PHRASING_NOTE}
{FORBIDDEN_PHRASES}
{NO_PASSIVE_LANGUAGE_NOTE}

Use the following example as a style guide only. DO NOT copy content or facts.

Example:
{INTRO_EXAMPLE}

Now write the Introduction section of a confidential mediation memorandum for {client_name}. Focus on tone, structure, and legal fluency. Use only the information provided below:

{input_text}
"""
    text = generate_with_openai(prompt)
    text = re.sub(r"\.\s+", ". ", text)
    return polish_text_for_legal_memo(text)

# === Mediation Memo Section Generators ===


PLAINTIFF_STATEMENT_EXAMPLE = """
On August 8, 2018, {{Plaintiff}}, then 38, was driving a commercial truck in the right lane of I-65 when his vehicle was rear-ended at a speed exceeding 50mph. Mr. Doe sustained life altering injuries to his neck and back. Mr. Doe suffered disc herniations at L3-4, L4-5, and L5-S1.  These injuries resulted in fiber tears, disc bulges, nerve injuries, radiating pain, dizziness, and limited mobility.   As a result of these injuries, he has undergone a variety of treatment, including, but not limited to physical therapy, injections, chiropractic treatment, electrical shock therapy, and prescription pain medication.
At the time of collision, Mr. Doe worked as a licensed commercial truck driver. (Ex. B, Doe Dep. 68-69). He had been driving commercially for four years since receiving his commercial driver’s license in 2014 (Ex. B, Doe Dep. 15). He had purchased part of his commercial truck with the goal of eventually starting his own commercial trucking company (Ex. B, Doe Dep. 88). Since the day of the crash, Mr. Doe has been unable to work as a commercial truck driver. (Ex. B, Doe Dep. 66). He now works as a dispatcher. Mr. Doe continues to be treated for neck and back pain. He will continue to require physical therapy, pain medication, and injections.
"""


def generate_plaintiff_statement(bio, client_name):
    prompt = f"""
{NO_HALLUCINATION_NOTE}
{LEGAL_FLUENCY_NOTE}

You are drafting a formal narrative background for a plaintiff in a legal memo. Use the following example only for style and tone. Do not copy or reuse any facts from the example.

---

💡 Example Paragraph (for tone only):
{PLAINTIFF_STATEMENT_EXAMPLE}

---

🎯 Task: Write a single-paragraph narrative background for the Plaintiff named {client_name}, using only the factual information provided below. Do not invent or assume facts. Start the paragraph with: "{client_name} is..."

📄 Provided Facts:
{bio}
"""
    text = generate_with_openai(prompt)
    text = re.sub(r"\.\s+", ". ", text)
    return polish_text_for_legal_memo(text)


DEFENDANT_STATEMENT_EXAMPLE = """

On August 8, 2018, {{Defendant1}}, then 43, was driving a commercial truck for the first time in the right lane of I-65 when he ran through Mr. Efimov’s truck. The police who responded to the accident cited Defendant Rakhimdjanov with following too closely behind Mr. Doe (Ex. C, Indiana Crash Report). After colliding into Mr. Efimov, Mr. Rakhimdjanov veered into the left, crashing into another vehicle, which crashed into the vehicle in front of it, before coming to a stop in the left lane (Ex. D, Rakhimdjanov Dep. 53).
At the time of the collision, Mr. Rakhimdjanov had been driving commercially for less than one day.  This was his first commercial driving job (Ex. D, Rakhimdjanov Dep. 30). Mr. Rakhimdjanov used an interpreter when he gave his deposition on December 16, 2021 and when he signed his employment contracts with STL Truckers. (Ex. D, Rakhimdjanov Dep. 73). However, Mr. Rakhimdjanov did not have an interpreter during his training with STL Truckers. (Ex. D, Rakhimdjanov Dep. 73). It is clear from the documents from STL that he did not even write his own name.
"""

def generate_defendant_statement(context, defendant_name):
    prompt = f"""
{NO_HALLUCINATION_NOTE}
{LEGAL_FLUENCY_NOTE}

You are drafting a formal narrative background for a defendant in a legal memo. Use the example below only as a reference for tone and structure. Do not reuse or reference any facts from the example.

---

💡 Example Paragraph (for tone only):
{DEFENDANT_STATEMENT_EXAMPLE}

---

🎯 Task: Write a single-paragraph narrative background for the Defendant named {defendant_name}, using only the provided factual details. Do not invent or assume any facts. Begin the paragraph with: "{defendant_name} is..."

📄 Provided Facts:
{context}
"""
    text = generate_with_openai(prompt)
    text = re.sub(r"\.\s+", ". ", text)
    return polish_text_for_legal_memo(text)



DEMAND_EXAMPLE = """
STL Trucking has represented to Plaintiff’s counsel that they only have a $1 million policy available for Mr. Efimov’s losses. To date, STL has yet to sign an affidavit verifying coverage. (Ex. F, Affidavit of No Excess Coverage).
"""

def generate_demand_section(summary, client_name):
    prompt = f"""
{NO_HALLUCINATION_NOTE}
{LEGAL_FLUENCY_NOTE}
{BAN_PHRASING_NOTE}
{FORBIDDEN_PHRASES}
{NO_PASSIVE_LANGUAGE_NOTE}

Write a mediation demand **paragraph** using the tone and clarity of the example below. 
This should be professional and persuasive—not a letter. Do not address “Dear Counsel” or sign off. This should simply be a summary. Keep it to two sentences. 

Example:
{DEMAND_EXAMPLE}

Facts:
{summary}
"""
    text = generate_with_openai(prompt)
    text = re.sub(r"\.\s+", ". ", text)
    return polish_text_for_legal_memo(text)



FACTS_LIABILITY_EXAMPLE = """
In the early afternoon of August 8, 2018, {{Plaintiff}} was operating a commercial semi-tractor trailer on I-65 near the 184-mile mark (Ex. C, Indiana Crash Report). It was a dry, sunny day, and Mr. Doe noted that first responders were at the site of a collision approximately one-quarter mile ahead (Ex. B, Doe Dep. 28-29). Traffic was at a complete stop (Ex. B, Doe Dep. 28). In accordance with his occupational training and four years of commercial truck driving experience, Mr. Doe slowly decelerated his vehicle and came to a stop two cars’ lengths—roughly 15 feet—behind the car in front of him (Ex. A, Doe Dep. 28, 35).
At the same time, defendant {{Defendant1}} was operating a commercial semi-tractor trailer for his first full day on the job. (Ex. D, Rakhimdjanov Dep. 35). Roughly 10 to 15 seconds after Mr. Doe had come to a complete stop, at approximately 1:45 PM, defendant violently crashed into the back of Mr. Efimov’s trailer (Ex. B, Doe Dep. 29). Defendant was traveling at a speed of about 60 to 65 miles per hour (Ex. C, Indiana Crash Report). The Defendant was traveling so fast at the time of collision that no tire marks were found at the scene.
Defendant driver’s truck knocked Stan’s truck off its frame.
The force of the collision caused Mr. Efimov’s head to violently snap forwards and backwards (Ex. A, Doe Dep. 86). Mr. Doe heard no warning sounds, no horns, and no break noises prior to the collision, and later found no skid marks on the road (Ex. A, Doe Dep. 39). In fact, defendant’s vehicle had so much momentum at the point of impact that defendant’s vehicle displaced Mr. Efimov’s commercial semi-tractor nearly 15 feet forward, causing the cab to come off its frame (Ex. A, Doe Dep. 41).
After making contact with Mr. Efimov’s truck, defendant veered into the left lane and collided with another vehicle (Ex. B, Indiana Crash Report). This vehicle was pushed forward, striking the vehicle in front of it, eventually coming to rest in the median (Ex. B, Indiana Crash Report). Three of the four vehicles involved in this incident had to be towed away (Ex. B, Indiana Crash Report). That day, Mr. Doe began experiencing neck pain, back pain, and dizziness (Ex. A, Doe Dep. 49). At his hotel in the morning, his pain was so extreme that he had to call his father to physically help him get out of bed, to go to the restroom, and to get into his car. He sought medical the next day back in Chicago (Ex. A, Doe Dep. 49).
"""
PARTIES_EXAMPLE = """
Plaintiff Stan Efimov is a resident of Cook County, Illinois. Defendant STL Truckers is a commercial carrier incorporated in Indiana, conducting business in Illinois, and was the employer of Defendant driver Arsen Rakhimdjanov. Defendant Arsen Rakhimdjanov was operating a commercial truck within the scope of his employment with STL Truckers at the time of the collision.
"""


def generate_party_section(party_details):
    prompt = f"""
{NO_HALLUCINATION_NOTE}
{LEGAL_FLUENCY_NOTE}

Write a “Parties” section using only the information below. Describe the Plaintiff and all Defendants, including any employer or corporate affiliations. Keep it factual, professional, and concise.

Example:
{PARTIES_EXAMPLE}

Input:
{party_details}
"""
    text = generate_with_openai(prompt)
    text = re.sub(r"\.\s+", ". ", text)
    return polish_text_for_legal_memo(text)

def generate_party_summary(plaintiff_names, defendant_names):
    """
    Generates a one-paragraph narrative describing all plaintiffs and all defendants.
    """
    plaintiff_list = ", ".join(plaintiff_names)
    defendant_list = ", ".join(defendant_names)

    prompt = f"""
{NO_HALLUCINATION_NOTE}
{LEGAL_FLUENCY_NOTE}

Write a single narrative paragraph describing all parties to the action.
Include each Plaintiff’s name, role, and jurisdictional detail. Do the same for all Defendants.
Use only this list:

Plaintiffs: {plaintiff_list}
Defendants: {defendant_list}

Keep it factual, formal, and brief — like the first paragraph of a complaint. Do not summarize allegations or damages here.
"""
    text = generate_with_openai(prompt)
    text = re.sub(r"\.\s+", ". ", text)
    return polish_text_for_legal_memo(text)


def generate_facts_liability_section(facts, deposition_text=None):
    prompt = f"""
{NO_HALLUCINATION_NOTE}
{LEGAL_FLUENCY_NOTE}

Draft the Facts / Liability section using only this information:
{facts}
"""
    if deposition_text:
        prompt += f"""
Carefully embed relevant direct quotes from the following deposition excerpts directly into the narrative as formal citations. Use phrasing like:
- “Plaintiff testified, ‘...’ (Ex. A, Efimov Dep. 43).”
- “According to deposition testimony, ‘...’ (Ex. A, Efimov Dep. 21).”

Only insert quotes where they support specific facts or legal reasoning.

Deposition excerpts for liability:
{deposition_text}
"""

    prompt += f"\n\nExample:\n{FACTS_LIABILITY_EXAMPLE}"
    text = generate_with_openai(prompt)
    text = re.sub(r"\.\s+", ". ", text)
    return polish_text_for_legal_memo(text)

CAUSATION_EXAMPLE = """
As a result of this occurrence, Stan will require a decompressive hemilaminectomy and microdiscectomy at L5-S1.
On August 8, 2018, Mr. {{Plaintiff}} was a restrained driver of a semi-tractor trailer stopped in traffic and rear-ended by another semi truck. On August 9, 2018 Mr. Doe saw Dr. Jaroslav Goldman at East West Internal Medicine Associates in Wheeling, IL to address pain he was experiencing as a result of the accident the day before. He complained of back pain, dizziness, fatigue, headaches, insomnia and neck pain. His assessment noted an acceleration-deceleration injury of the neck, muscle spasms and external constriction of the neck. He was prescribed to begin physical therapy, chiropractic manipulations and X-rays of the lumbosacral spine. (Ex. G, Doe Medical Records).

Mr. Doe began chiropractic treatment with Dr. Kaspars Vilems, DC on August 10, 2018. The treatment plan was intended to address Mr. Efimov’s pain in his neck and upper and lower back. Dr. Vilems performed a series of chiropractic manipulations, electric stimulations, and therapeutic exercises at each appointment (Ex. G, Doe Medical Records). Mr. Doe continued this therapy until October 19, 2018. At that point he had attended 21 chiropractic therapy visits.
On August 28, 2018, Mr. Doe underwent an MRI of the lumbar spine ordered by Dr. Vilems for his low back pain. The MRI revealed a high-density zone in the left lateral fibers of the L4-5 disc suggestive of an annular tear.
On October 4, 2018, Mr. Doe had a pain consultation with Dr. Yuriy Bukhalo at Northwest Suburban Pain Center for his bilateral low back pain radiating to his right knee. Mr. Doe reported that he began feeling pain after sitting for more than 15 minutes, making it difficult to continue working as a truck driver. Dr. Bukhalo agreed that the annular tear detected in the lumbar spine MRI was likely the cause of this pain and recommended intensifying physical therapy, wearing a brace, and initiating an anti-inflammatory.
At a follow-up appointment on October 23, Dr. Bukhalo performed right L4-5 and L5-S1 transforaminal epidural steroid injections (TFESIs). At the next appointment on November 6, Mr. Doe reported 60% ongoing pain improvement. He still felt pain while sitting for long periods of time. Due to this pain, he was forced to change his job as a truck driver to a managerial position with the trucking company he drove for (Ex. A, Doe Dep. 67). Dr. Bukhalo performed the same TFESIs at this appointment (Ex. G, Doe Medical Records).

By November 12, 2019, Mr. Efimov’s lower back pain had not subsided. He had a surgical consultation with Dr. Sean Salehi at the Neurological Surgery & Spine Surgery S.C. to address his continuing low back pain. Dr. Salehi found that Mr. Doe was not a surgical candidate due to his elevated BMI and intermittent symptoms and was referred to pain management instead.
Mr. Doe scheduled an appointment for December 4, 2019, with Dr. Krishna Chunduri at Advanced Spine and Pain Specialists for his lower back pain. He was prescribed a Medrol Dosepak for his pain flare-ups.
On March 17, 2020, he had a visit with Dr. Mark Farag at Midwest Anesthesia and Pain Specialists. Mr. Doe reported that the injections performed by Dr. Bukhalo took the pain away for approximately three to four months. The assessment noted low back pain, lumbar radiculopathy, lumbar intervertebral disc displacement, and lumbosacral spondylosis with myelopathy or radiculopathy. Dr. Farag recommended an L5-S1 lumbar epidural steroid injection (LESI).
Mr. Doe followed up with Dr. Farag on January 26, 2021. He reported worsening back pain with intermittent numbness and tingling in the right lower extremity. At this point, he had completed physical therapy and was participating in a home exercise program. Dr. Farag once again recommended an L5-S1 LESI. This injection was performed on February 18, 2021.
On March 4, 2021, Mr. Doe reported to Dr. Farag that he had an improvement in pain since his injection, but still felt intermittent radiating pain in his right leg. Dr. Farag ordered right L4-L5 and L5-S1 TFESIs. These injections were performed on May 27, 2021.
His most recent MRI occurred on June 20, 2022 at Empire Imaging in Pembroke Pines, FL. This scan found a disc bulge with a superimposed right foraminal disc herniation at L3-4 and a disc bulge with a left foraminal disc herniation at L4-5. At L5-S1 there is a disc bulge with superimposed right posterior disc herniation.
At present, Mr. Doe reports that his conditions remain symptomatic and cause ongoing disability including sciatic nerve pain. He reports that his low back and lower extremity pain remains present despite undergoing extensive therapy.

"""
def generate_causation_injuries(causation_info):
    prompt = f"""
{NO_HALLUCINATION_NOTE}
{LEGAL_FLUENCY_NOTE}

Write a paragraph connecting the incident to the client’s injuries and medical course.

Example:
{CAUSATION_EXAMPLE}

Facts:
{causation_info}
"""
    text = generate_with_openai(prompt)
    text = re.sub(r"\.\s+", ". ", text)
    return polish_text_for_legal_memo(text)


HARMS_EXAMPLE = """
Prior to this collision, Mr. Doe was a healthy and active 38-year-old man who enjoyed playing basketball, tennis, and soccer at the park with his friends (Ex. A, Doe Dep. 77). 
Since the day of the crash, Stan has not been able to drive a truck.  He physically cannot climb into a truck, sit for long hours, work to unhitch the trailer, and climb out of the truck.  He enjoyed playing with and lifting up his niece, who weighs 25 pounds (Ex. A, Doe Dep. 79). Since the accident, he has been unable to partake in these activities that once brought him great joy. Mr. Doe had been driving commercial trucks since 2014 and was forced to step away from driving because of the intense pain he experienced after sitting for prolonged periods of time (Ex. A, Doe Dep. 67). As a result of the collision and to this date, he has been unable to do any these activities without experiencing severe, debilitating pain and has, therefore, lost the quality of life he once had. He had no previous neck or back injuries before this accident.
"""
FUTURE_BILLS_EXAMPLE = """
Mr. Doe’s ongoing symptoms will require lifelong pain management. Based on his treating providers’ evaluations and the history of pain persistence despite conservative treatment, Plaintiff is expected to undergo future interventions such as repeat TFESI injections and potentially spinal surgery. The cost of these procedures, including follow-up physical therapy, imaging, and medication, is estimated to exceed $150,000 over his lifetime.
"""

def generate_additional_harms(harm_info, deposition_text=None):
    prompt = f"""
{NO_HALLUCINATION_NOTE}
{LEGAL_FLUENCY_NOTE}

Write the “Additional Harms and Losses” section based on the information below.

Input:
{harm_info}
"""
    if deposition_text:
        prompt += f"""

Carefully embed relevant direct quotes from the following deposition excerpts directly into the narrative as formal citations. Use phrasing like:
- “Plaintiff testified, ‘...’ (Ex. A, Efimov Dep. 43).”
- “According to deposition testimony, ‘...’ (Ex. A, Efimov Dep. 21).”

Only insert quotes where they support specific facts or legal reasoning.
{deposition_text}
"""

    prompt += f"\n\nExample:\n{HARMS_EXAMPLE}"
    text = generate_with_openai(prompt)
    text = re.sub(r"\.\s+", ". ", text)
    return polish_text_for_legal_memo(text)


def generate_future_medical(future_info, deposition_text=None):
    prompt = f"""
{NO_HALLUCINATION_NOTE}
{LEGAL_FLUENCY_NOTE}

Draft a future medical expenses section using this tone:

{FUTURE_BILLS_EXAMPLE}

Input:
{future_info}
"""
    if deposition_text:
        prompt += f"""

You may quote directly from the following deposition excerpts to reinforce impact on daily life, work, or emotional well-being:

Deposition excerpts for damages:
{deposition_text}
"""
    text = generate_with_openai(prompt)
    text = re.sub(r"\.\s+", ". ", text)
    return polish_text_for_legal_memo(text)


CONCLUSION_EXAMPLE = """
Plaintiff’s past and future medical bills alone nearly exceed the purported $1 million dollar 
policy that STL Truckers has for their and their Driver’s punitive actions causing Stan life altering injuries and damages. Stan, now 42, has 30 plus years of pain, suffering, and loss of normal life related to this occurrence caused by STL Truckers putting a trucker on the road who was not trained for the job.
"""


def generate_conclusion_section(notes):
    prompt = f"""
{NO_HALLUCINATION_NOTE}
{LEGAL_FLUENCY_NOTE}

Write the closing section of a mediation memo.

Example:
{CONCLUSION_EXAMPLE}

Input:
{notes}
"""
    text = generate_with_openai(prompt)
    text = re.sub(r"\.\s+", ". ", text)
    return polish_text_for_legal_memo(text)


# === Improved Placeholder Replacer ===
def replace_placeholders(doc, replacements):
    def rebuild_paragraph(paragraph):
        combined_text = "".join(run.text for run in paragraph.runs)
        for key, val in replacements.items():
            combined_text = combined_text.replace(key, val)
        if combined_text.strip():  # Only replace if there's something to write
            # Clear and reset
            for run in paragraph.runs:
                run.text = ""

            if paragraph.runs:
                paragraph.runs[0].text = combined_text
            else:
                paragraph.add_run(combined_text)


            if not paragraph.runs:
                paragraph.add_run()

    def replace_in_cell(cell: _Cell):
        for paragraph in cell.paragraphs:
            rebuild_paragraph(paragraph)

    for paragraph in doc.paragraphs:
        rebuild_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell)

# === Template Filler ===

def fill_mediation_template(data, template_path, output_path):
    from docx import Document
    import os
    from datetime import datetime

    doc = Document(template_path)

    # Start with updated static placeholders
    replacements = {
        "{{Court}}": data.get("court", ""),
        "{{Case_Number}}": data.get("case_number", ""),
        "{{Introduction}}": data.get("introduction", ""),
        "{{Parties}}": data.get("parties", ""),
        "{{Demand}}": data.get("demand", ""),
        "{{Facts_Liability}}": data.get("facts_liability", ""),
        "{{Causation_Injuries_Treatment}}": data.get("causation_injuries", ""),
        "{{Additional_Harms_Losses}}": data.get("additional_harms", ""),
        "{{Future_Medical_Bills}}": data.get("future_bills", ""),
        "{{Conclusion}}": data.get("conclusion", ""),
        "{{Plaintiffs}}": data.get("Plaintiffs", ""),
        "{{Defendants}}": data.get("Defendants", "")
    }

    for i in range(1, 4):
        name = data.get(f"plaintiff{i}", "")
        statement = data.get(f"plaintiff{i}_statement", "")
        combined = f"**{name}**\n\n{statement}" if name or statement else ""
        replacements[f"{{{{Plaintiff_{i}}}}}"] = combined


    for i in range(1, 8):
        name = data.get(f"defendant{i}", "")
        statement = data.get(f"defendant{i}_statement", "")
        combined = f"**{name}**\n\n{statement}" if name or statement else ""
        replacements[f"{{{{defendant_{i}}}}}"] = combined


    def rebuild_paragraph(paragraph):
        original_text = "".join(run.text for run in paragraph.runs)
        replaced_text = original_text
        for key, val in replacements.items():
            replaced_text = replaced_text.replace(key, val)

        if replaced_text != original_text:
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = replaced_text
            else:
                paragraph.add_run(replaced_text)

    def replace_in_cell(cell):
        for paragraph in cell.paragraphs:
            rebuild_paragraph(paragraph)

    for paragraph in doc.paragraphs:
        rebuild_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell)

    plaintiff_name = data.get("plaintiff", "Unknown").replace(" ", "_")
    filename = f"Mediation_Memo_{plaintiff_name}_{datetime.today().strftime('%Y-%m-%d')}.docx"
    output_file_path = os.path.join(output_path, filename)
    print("=== PLACEHOLDERS SENT TO TEMPLATE ===")
    for k, v in replacements.items():
        if '{{' in k:
            print(f"{k}: {'[FILLED]' if v.strip() else '[EMPTY]'}")

    doc.save(output_file_path)

    return output_file_path


# --- Safe wrapper to handle OpenAI rate limits ---

import re

def redact_text(text):
    """
    Redacts sensitive info (PHI) from already OCR'd or typed text.
    """
    text = re.sub(
        r"\b(?:DOB|D\.O\.B)\s*[:\-]?\s*\d{1,2}[\/\-]\d{1,2}[\/\-]\d{2,4}", "[REDACTED DOB]", text, flags=re.I
    )
    text = re.sub(
        r"\b\d{3}-\d{2}-\d{4}\b", "[REDACTED SSN]", text
    )
    text = re.sub(
        r"\b(Name|Patient)\s*[:\-]?\s*[A-Z][a-z]+\s[A-Z][a-z]+", "[REDACTED NAME]", text
    )
    return text

# --- Main time split function ---
def generate_quotes_in_chunks(text_chunks, delay_seconds=10):
    """
    Categorize Q&A deposition quotes by Liability and Damages.
    Ensures line numbers, formatting, and strict legal output.
    """
    liability_quotes = []
    damages_quotes = []

    for i, chunk in enumerate(text_chunks):
        sub_chunks = chunk_text(chunk, max_chars=6000)  # Safely re-chunk each chunk

        for j, sub_chunk in enumerate(sub_chunks):
            prompt = f"""
You are a litigation analyst. Categorize the following deposition excerpts into **Liability** or **Damages**. 

🧾 **Return Format (strict)**:
Only include bullet points like this:
- **0012:25 Q:** "What were you responsible for on that day?"  
  **0012:26 A:** "I was supervising the road closure."

📂 **Categories**:
**Liability** = Questions or answers about duties, fault, conduct, knowledge of events, observations, or cause of the incident.  
**Damages** = Anything about pain, treatment, limitations, daily impact, job loss, suffering, or future care.

⚠️ **Rules**:
- Each bullet must include both the Q and A with proper line numbers.
- Quotes must be exact. No paraphrasing. No summaries. No interpretation.
- Use quotation marks. Do not reformat Qs or As.
- If a line is cut off, include only the first full line.
- No duplicates. Do not repeat quotes across categories or chunks.

📄 **Excerpt**:
{sub_chunk}
"""

            try:
                result = safe_generate(generate_with_openai, prompt)

                # Parse out the quotes
                liability_section = re.findall(r"\*\*Liability\*\*(.*?)(?=\*\*Damages\*\*|$)", result, re.DOTALL)
                damages_section = re.findall(r"\*\*Damages\*\*(.*)", result, re.DOTALL)

                if liability_section:
                    liability_quotes.append(liability_section[0].strip())
                if damages_section:
                    damages_quotes.append(damages_section[0].strip())

                print(f"Processed chunk {i+1}.{j+1}/{len(text_chunks)}")
            except APIStatusError as e:
                print(f"API error on chunk {i+1}.{j+1}: {e}")
                raise e

            if j < len(sub_chunks) - 1:
                time.sleep(delay_seconds)

    def clean_and_dedup(quotes):
        combined = "\n".join(quotes)
        seen = set()
        cleaned = []
        for line in combined.splitlines():
            line = line.strip()
            if line and line not in seen:
                seen.add(line)
                cleaned.append(line)
        return "\n".join(cleaned)

    return {
        "liability_quotes": clean_and_dedup(liability_quotes),
        "damages_quotes": clean_and_dedup(damages_quotes)
    }

def split_and_combine(fn, long_text, quotes="", chunk_size=3000):
    chunks = [long_text[i:i+chunk_size] for i in range(0, len(long_text), chunk_size)]
    results = []
    for chunk in chunks:
        if quotes:
            results.append(safe_generate(fn, chunk, quotes))
        else:
            results.append(safe_generate(fn, chunk))
    return "\n\n".join(results)
    # --- Main generation function ---

def polish_text_for_legal_memo(text):
    prompt = f"""
You are a senior legal editor. Improve the following legal writing by:
- Converting to active voice
- Eliminating redundant or vague phrases
- Removing repetition of facts already stated
- Keeping it formal, persuasive, and professionally polished
- Do not introduce new facts or change case theory

Text:
{text}
"""
    return safe_generate(generate_with_openai, prompt)

def generate_memo_from_summary(data, template_path, output_dir, text_chunks):
    import time
    memo_data = {}

    memo_data["Plaintiffs"] = ", ".join(
        [data.get(f"plaintiff{i}", "") for i in range(1, 4) if data.get(f"plaintiff{i}", "")]
    )
    memo_data["Defendants"] = ", ".join(
        [data.get(f"defendant{i}", "") for i in range(1, 8) if data.get(f"defendant{i}", "")]
    )

    memo_data["court"] = data["court"]
    memo_data["case_number"] = data["case_number"]

    # Determine primary plaintiff
    plaintiff1 = data.get("plaintiff1") or data.get("plaintiff") or "Plaintiff"
    memo_data["plaintiff"] = plaintiff1
    memo_data["plaintiff1"] = plaintiff1
    memo_data["plaintiff1_statement"] = safe_generate(
        generate_plaintiff_statement, trim_to_token_limit(data["complaint_narrative"], 4000), plaintiff1
    )
    time.sleep(20)

    # === Handle Plaintiffs ===
    for i in range(1, 4):
        name = data.get(f"plaintiff{i}", "").strip()
        if name:
            party_input = (
                trim_to_token_limit(data.get("party_information_from_complaint", ""), 3000)
                + "\n\n"
                + trim_to_token_limit(data.get("settlement_summary", ""), 2000)
            )
            statement = safe_generate(generate_plaintiff_statement, party_input, name)
        else:
            statement = ""

        memo_data[f"plaintiff{i}"] = name
        memo_data[f"plaintiff{i}_statement"] = statement

    # === Handle Defendants ===
    for i in range(1, 8):
        name = data.get(f"defendant{i}", "").strip()
        if name:
            party_input = (
                trim_to_token_limit(data.get("party_information_from_complaint", ""), 3000)
                + "\n\n"
                + trim_to_token_limit(data.get("settlement_summary", ""), 2000)
            )
            statement = safe_generate(generate_defendant_statement, party_input, name)
        else:
            statement = ""

        memo_data[f"defendant{i}"] = name
        memo_data[f"defendant{i}_statement"] = statement


 # Main body content
    memo_data["introduction"] = polish_text_for_legal_memo(
        safe_generate(generate_introduction, trim_to_token_limit(data["complaint_narrative"], 4000), plaintiff1)
    )
    time.sleep(20)

    memo_data["demand"] = polish_text_for_legal_memo(
        safe_generate(generate_demand_section, trim_to_token_limit(data["settlement_summary"], 2000), plaintiff1)
    )
    time.sleep(20)

    quotes_dict = generate_quotes_in_chunks(text_chunks, delay_seconds=20)
    trimmed_medical_summary = trim_to_token_limit(data.get("medical_summary", ""), 10000)
    liability_quotes = quotes_dict["liability_quotes"]
    damages_quotes = quotes_dict["damages_quotes"]

    memo_data["facts_liability"] = polish_text_for_legal_memo(
        safe_generate(generate_facts_liability_section, 
        "\n\n".join(chunk_text(data["complaint_narrative"])), 
            liability_quotes
        )
    )

    time.sleep(20)

    memo_data["additional_harms"] = polish_text_for_legal_memo(
        safe_generate(generate_additional_harms, trimmed_medical_summary, damages_quotes)
    )

    time.sleep(20)

    memo_data["future_bills"] = polish_text_for_legal_memo(
        "\n\n".join([
            safe_generate(generate_future_medical, chunk, trim_to_token_limit(damages_quotes, 2000))
            for chunk in chunk_text(trimmed_medical_summary)
        ])
    )
    time.sleep(20)

    memo_data["causation_injuries"] = polish_text_for_legal_memo(
        "\n\n".join([
            safe_generate(generate_causation_injuries, chunk)
            for chunk in chunk_text(trimmed_medical_summary)
        ])
    )
    time.sleep(20)

# === FORMAL NARRATIVE PARTIES SECTION WITH HEADINGS ===
    plaintiff_sections = []
    defendant_sections = []

    # === Collect narrative paragraphs for display under "Parties" ===
    for i in range(1, 4):
        name = memo_data.get(f"plaintiff{i}", "")
        statement = memo_data.get(f"plaintiff{i}_statement", "")
        if name or statement:
            plaintiff_sections.append(f"Plaintiff {name} Statement:\n{statement}")

    for i in range(1, 8):
        name = memo_data.get(f"defendant{i}", "")
        statement = memo_data.get(f"defendant{i}_statement", "")
        if name or statement:
            defendant_sections.append(f"Defendant {name} Statement:\n{statement}")


    # === Generate narrative and full parties section ===
    plaintiff_names = [memo_data.get(f"plaintiff{i}", "") for i in range(1, 4) if memo_data.get(f"plaintiff{i}", "")]
    defendant_names = [memo_data.get(f"defendant{i}", "") for i in range(1, 8) if memo_data.get(f"defendant{i}", "")]

    memo_data["parties"] = polish_text_for_legal_memo(
        safe_generate(generate_party_summary, plaintiff_names, defendant_names)
    )

    if plaintiff_sections:
        memo_data["parties"] += "\n\nPLAINTIFFS:\n" + "\n\n".join(plaintiff_sections) + "\n\n"
    if defendant_sections:
        memo_data["parties"] += "DEFENDANTS:\n" + "\n\n".join(defendant_sections)


    # === Final cleanup and formatting ===
    memo_data["conclusion"] = polish_text_for_legal_memo(
        safe_generate(generate_conclusion_section, data["settlement_summary"])
    )

    intro_clean = memo_data.pop("introduction", "")
    facts_text = data["complaint_narrative"]
    memo_data["Introduction"] = polish_text_for_legal_memo(intro_clean.replace(facts_text[:200], ""))
    memo_data["Demand"] = memo_data.pop("demand", "")
    memo_data["Facts_Liability"] = memo_data.pop("facts_liability", "")
    memo_data["Causation_Injuries_Treatment"] = memo_data.pop("causation_injuries", "")
    memo_data["Additional_Harms_Losses"] = memo_data.pop("additional_harms", "")
    memo_data["Future_Medical_Bills"] = memo_data.pop("future_bills", "")
    memo_data["Conclusion"] = memo_data.pop("conclusion", "")

    for key in ["Introduction", "Facts_Liability", "Additional_Harms_Losses", "Future_Medical_Bills", "Conclusion"]:
        if memo_data.get(key):
            memo_data[key] = re.sub(r"\s{2,}", " ", memo_data[key].strip())

    file_path = fill_mediation_template(memo_data, template_path, output_dir)
    return file_path, memo_data

def generate_plaintext_memo(memo_data):
    """
    Fallback: Generate a plain text version of the memo from all memo_data.
    """
    sections = [
        ("Court", memo_data.get("court", "")),
        ("Case Number", memo_data.get("case_number", "")),
        ("Introduction", memo_data.get("introduction", "")),
        ("Parties", memo_data.get("parties", "")),
        ("Demand", memo_data.get("demand", "")),
        ("Facts / Liability", memo_data.get("facts_liability", "")),
        ("Causation, Injuries, and Treatment", memo_data.get("causation_injuries", "")),
        ("Additional Harms and Losses", memo_data.get("additional_harms", "")),
        ("Future Medical Bills", memo_data.get("future_bills", "")),
        ("Conclusion", memo_data.get("conclusion", ""))
    ]

    # Add plaintiff and defendant blocks
    for i in range(1, 4):
        name = memo_data.get(f"plaintiff{i}", "")
        statement = memo_data.get(f"plaintiff{i}_statement", "")
        if name or statement:
            sections.append((f"Plaintiff {i}: {name}", statement))

    for i in range(1, 8):
        name = memo_data.get(f"defendant{i}", "")
        statement = memo_data.get(f"defendant{i}_statement", "")
        if name or statement:
            sections.append((f"Defendant {i}: {name}", statement))

    # Build plain text
    lines = []
    for title, content in sections:
        if content.strip():
            lines.append(f"=== {title.upper()} ===\n{content.strip()}\n")

    return "\n".join(lines)