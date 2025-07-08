from docx import Document
from datetime import datetime
import os

def fill_template(template_path, output_dir, replacements):
    # === Load Template ===
    try:
        doc = Document(template_path)
    except Exception as e:
        print(f"❌ Failed to load template: {e}")
        return

    def rebuild_paragraph(paragraph):
        original_text = "".join(run.text for run in paragraph.runs)
        replaced_text = original_text
        for key, val in replacements.items():
            replaced_text = replaced_text.replace(key, val)

        if replaced_text != original_text:
            for run in paragraph.runs:
                run.text = ""
            if paragraph.runs:
                paragraph.runs[0].text = replaced_text
            else:
                paragraph.add_run(replaced_text)

    # === Replace in Paragraphs ===
    for paragraph in doc.paragraphs:
        rebuild_paragraph(paragraph)

    # === Replace in Tables ===
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    rebuild_paragraph(paragraph)

    # === Prepare Output Directory ===
    os.makedirs(output_dir, exist_ok=True)

    # === Save Document ===
    filename = f"Mediation_Memo_Test_{datetime.today().strftime('%Y-%m-%d')}.docx"
    output_path = os.path.join(output_dir, filename)

    print("\n=== PLACEHOLDER STATUS ===")
    for k, v in replacements.items():
        if "{{" in k:
            status = "[FILLED]" if v.strip() else "[EMPTY]"
            print(f"{k}: {status}")

    try:
        doc.save(output_path)
        print(f"\n✅ Memo saved to: {output_path}")
    except Exception as e:
        print(f"❌ Failed to save document: {e}")

# === TEST DATA ===
replacements = {
    "{{Court}}": "Circuit Court of Cook County",
    "{{Case_Number}}": "2025-L-123456",
    "{{Plaintiff_1_Name}}": "John Doe",
    "{{Plaintiff_1_Statement}}": "John Doe is a long-haul truck driver who sustained permanent injuries...",
    "{{Plaintiff_2_Name}}": "",
    "{{Plaintiff_2_Statement}}": "",
    "{{Defendant_1_Name}}": "STL Trucking Co.",
    "{{Defendant_1_Statement}}": "STL Trucking employed the driver at fault and failed to train him properly...",
    "{{Demand}}": "$1,000,000 policy demand due to life-altering injuries.",
    "{{Facts_Liability}}": "On April 1, 2022, the defendant rear-ended the plaintiff on I-70.",
    "{{Causation_Injuries_Treatment}}": "Plaintiff required surgery and 18 months of physical therapy.",
    "{{Additional_Harms_Losses}}": "Plaintiff can no longer drive professionally.",
    "{{Future_Medical_Bills}}": "Expected costs exceed $200,000.",
    "{{Conclusion}}": "We are prepared to proceed to trial if resolution is not reached."
}

# === RUN ===
fill_template(r"C:\LegalAutomationHub - v2\templates\mediation_template.docx", "outputs", replacements)

