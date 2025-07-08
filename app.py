import streamlit as st
st.set_page_config(page_title="Legal Automation Hub", layout="wide")

from scripts.run_foia import run_foia
from scripts.run_demand import run
from scripts.run_mediation import (
    generate_with_openai,
    generate_introduction,
    generate_plaintiff_statement,
    generate_defendant_statement,
    generate_demand_section,
    generate_facts_liability_section,
    generate_causation_injuries,
    generate_additional_harms,
    generate_future_medical,
    generate_conclusion_section,
    generate_quotes_in_chunks,
    fill_mediation_template,
)
from scripts.run_mediation import split_and_combine
from scripts.run_mediation import generate_quotes_in_chunks


import pandas as pd
import os
import dropbox
from io import BytesIO
import zipfile
import io
import tempfile
from docx import Document
from datetime import datetime
from users import USERS, hash_password



import time
import openai
try:
    from openai.error import RateLimitError  # older SDKs
except ImportError:
    try:
        from openai.errors import RateLimitError  # newer SDKs
    except ImportError:
        RateLimitError = Exception  # fallback

def preview_first_page(template_path):
    """
    Loads the first page of a Word document and returns its text.
    """
    try:
        doc = Document(template_path)
        text_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        first_page_text = "\n".join(text_lines[:25])  # Adjust line count as needed
        return first_page_text
    except Exception as e:
        return f"❌ Error loading preview: {e}"


def safe_generate(func, *args, max_retries=3, delay=5, **kwargs):
    """
    Retry wrapper to safely call GPT-powered functions like generate_xxx with backoff on RateLimitError.
    """
    for attempt in range(max_retries):
        try:
            return func(*args, **kwargs)
        except RateLimitError:
            if attempt < max_retries - 1:
                time.sleep(delay * (attempt + 1))  # Exponential backoff
            else:
                raise
        except Exception as e:
            raise RuntimeError(f"Generation failed: {e}")

def extract_quotes_from_text(ocr_text, user_instructions=""):
    base_prompt = """
You are a legal assistant. Given this deposition or record text, extract the most relevant direct quotes (verbatim, in quotes) that support either:
- liability,
- injuries,
- or harm to quality of life.
"""
    if user_instructions.strip():
        base_prompt += f"\n\nThe user has provided the following additional instructions:\n{user_instructions.strip()}\n"

    full_prompt = f"""{base_prompt}

Only return a list of quotes. Do not paraphrase. Only use what is in the input.

Input:
{ocr_text}
"""
    return generate_with_openai(full_prompt)



# === Username + Password Login ===
if "logged_in" not in st.session_state:
    st.session_state.logged_in = False

if not st.session_state.logged_in:
    st.title("Login")

    username = st.text_input("Username")
    password = st.text_input("Password", type="password")

    if st.button("Login"):
        hashed_input = hash_password(password)
        if username in USERS and USERS[username] == hashed_input:
            st.session_state.logged_in = True
            st.session_state.username = username
            st.success(f"Welcome, {username}!")
            # ✅ No need for st.experimental_rerun()
        else:
            st.error("Invalid username or password")

    st.stop()
else:
    st.sidebar.markdown(f"**Logged in as:** `{st.session_state.username}`")

st.markdown("""
<style>
.stButton > button {
    background-color: #B08B48;
    color: white;
    border: none;
    padding: 0.5rem 1rem;
    font-weight: 600;
}
.stTextInput > div > input {
    border: 1px solid #0A1D3B;
}
.stTextArea > div > textarea {
    border: 1px solid #0A1D3B;
}
</style>
""", unsafe_allow_html=True)

# === Branding: Logo inside navy header bar ===
import base64

def load_logo_base64(file_path):
    with open(file_path, "rb") as image_file:
        return base64.b64encode(image_file.read()).decode()

logo_base64 = load_logo_base64("sggh_logo.png")

st.markdown(f"""
<div style="background-color: #0A1D3B; padding: 2rem 0; text-align: center;">
    <img src="data:image/png;base64,{logo_base64}" width="360" style="margin-bottom: 1rem;" />
    <h1 style="color: white; font-size: 2.2rem; margin: 0;">Stinar Gould Grieco & Hensley</h1>
</div>
""", unsafe_allow_html=True)

# === Sidebar Navigation ===
with st.sidebar:
    st.markdown("### ⚖️ Legal Automation Hub")
    tool = st.radio("Choose Tool", [
        "📖 Instructions & Support",
        "📄 Batch Doc Generator",
        "📬 FOIA Requests",
        "📂 Demands",
        "📊 Litigation Dashboard",
        "🧾 Mediation Memos",
        "🚧 Complaint (In Progress)",
        "🚧 Subpoenas (In Progress)",
    ])


# === FOIA Requests ===
if tool == "📬 FOIA Requests":
    st.header("📨 Generate FOIA Letters")

    with st.form("foia_form"):
        client_id = st.text_input("Client ID")
        defendant_name = st.text_input("Recipient Name")
        abbreviation = st.text_input("Recipient Abbreviation (for file name)")
        address_line1 = st.text_input("Recipient Address Line 1")
        address_line2 = st.text_input("Recipient Address Line 2 (City, State, Zip)")
        date_of_incident = st.date_input("Date of Incident")
        location = st.text_input("Location of Incident")
        case_synopsis = st.text_area("Case Synopsis")
        potential_requests = st.text_area("Potential Requests (can be reused from another)")
        explicit_instructions = st.text_area("Explicit Instructions (optional)")
        case_type = st.text_input("Case Type")
        facility = st.text_input("Facility or System")
        defendant_role = st.text_input("Recipient Role")

        submitted = st.form_submit_button("Generate FOIA Letter")

    if submitted:
        try:
            df = pd.DataFrame([{
                "Client ID": client_id,
                "Recipient Name": defendant_name,
                "Recipient Abbreviation": abbreviation,
                "Recipient Line 1 (address)": address_line1,
                "Recipient Line 2 (City,state, zip)": address_line2,
                "DOI": date_of_incident,
                "location of incident": location,
                "Case Synopsis": case_synopsis,
                "Potential Requests": potential_requests,
                "Explicit instructions": explicit_instructions,
                "Case Type": case_type,
                "Facility or System": facility,
                "Recipient Role": defendant_role
            }])

            output_paths = run_foia(df)
            st.success("✅ FOIA letter generated!")

            from docx import Document

            for path in output_paths:
                filename = os.path.basename(path)

                # Preview contents
                doc = Document(path)
                preview_text = "\n".join([p.text for p in doc.paragraphs if p.text.strip()])
                st.subheader(f"📄 {filename}")
                st.text_area("📘 Preview", preview_text, height=400)

                # Download button
                with open(path, "rb") as f:
                    st.download_button(
                        label=f"⬇️ Download {filename}",
                        data=f.read(),
                        file_name=filename,
                        mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                    )

        except Exception as e:
            st.error(f"❌ Error: {e}")

# === Demands ===
elif tool == "📂 Demands":
    st.header("📑 Generate Demand Letters")

    st.subheader("📋 Fill in Demand Letter Info")

    with st.form("demand_form"):
        client_name = st.text_input("Client Name")
        defendant = st.text_input("Defendant")
        incident_date = st.date_input("Incident Date")
        location = st.text_input("Location")
        summary = st.text_area("Summary of Incident")
        damages = st.text_area("Damages")

        submitted = st.form_submit_button("Generate Demand Letter")

    if submitted:
        import pandas as pd
        from datetime import datetime

        df = pd.DataFrame([{
            "Client Name": client_name,
            "Defendant": defendant,
            "IncidentDate": incident_date.strftime("%B %d, %Y"),
            "Location": location,
            "Summary": summary,
            "Damages": damages
        }])

        try:
            output_paths = run(df)
            st.success("✅ Letter generated!")

            for path in output_paths:
                filename = os.path.basename(path)
                with open(path, "rb") as f:
                    st.download_button(
                        label=f"Download {filename}",
                        data=f,
                        file_name=filename
                    )
        except Exception as e:
            st.error(f"❌ Error: {e}")

# === Routing ===
if tool == "📄 Batch Doc Generator":
    st.header("📄 Batch Document Generator")

    TEMPLATE_FOLDER = os.path.join("templates", "batch_docs")
    os.makedirs(TEMPLATE_FOLDER, exist_ok=True)

    try:
        campaign_df = pd.read_csv("campaigns.csv")
        CAMPAIGN_OPTIONS = sorted(campaign_df["Campaign"].dropna().unique())
    except Exception as e:
        CAMPAIGN_OPTIONS = []
        st.error(f"❌ Failed to load campaigns.csv: {e}")

    st.markdown("""
    > **How it works:**  
    > 1. Upload a template with `{placeholders}`  
    > 2. Upload Excel with matching column headers  
    > 3. Enter filename format, generate, and download

    ✅ Validates fields  
    📁 Version-safe template naming  
    🔐 No coding required
    """)

    st.subheader("📟 Template Manager")
    template_mode = st.radio("Choose an action:", ["Upload New Template", "Select a Saved Template", "Template Options"])

    def process_and_preview(template_path, df, output_name_format):
        # === Clean up data ===
        df = df.copy()

        # Format DOBs if present
        if "DOB" in df.columns:
            df["DOB"] = pd.to_datetime(df["DOB"], errors="coerce").dt.strftime("%m/%d/%Y")

        # Replace NaN in all columns with blank strings
        df = df.fillna("")

        st.subheader("🔍 Preview First Row of Excel Data")
        st.dataframe(df.head(1))

        st.markdown("**Columns in Excel:**")
        st.code(", ".join(df.columns))

        preview_filename = output_name_format
        for key, val in df.iloc[0].items():
            preview_filename = preview_filename.replace(f"{{{{{key}}}}}", str(val))
        st.markdown("**📄 Preview Filename for First Row:**")
        st.code(preview_filename)

        left, right = "{{", "}}"
        with tempfile.TemporaryDirectory() as temp_dir:
            word_dir = os.path.join(temp_dir, "Word Documents")
            os.makedirs(word_dir)

            for idx, row in df.iterrows():
                row = row.fillna("").to_dict()

                for k, v in row.items():
                    if isinstance(v, (pd.Timestamp, datetime)):
                        row[k] = v.strftime("%m/%d/%Y")

                doc = Document(template_path)

                for para in doc.paragraphs:
                    for key, val in row.items():
                        placeholder = f"{left}{key}{right}"
                        for run in para.runs:
                            if placeholder in run.text:
                                run.text = run.text.replace(placeholder, str(val))

                for table in doc.tables:
                    for cell in table._cells:
                        for para in cell.paragraphs:
                            for run in para.runs:
                                for key, val in row.items():
                                    placeholder = f"{left}{key}{right}"
                                    if placeholder in run.text:
                                        run.text = run.text.replace(placeholder, str(val))

                name_for_file = output_name_format
                for key, val in row.items():
                    name_for_file = name_for_file.replace(f"{left}{key}{right}", str(val))
                filename = name_for_file + ".docx"

                doc_path = os.path.join(word_dir, filename)
                doc.save(doc_path)

            zip_buffer = io.BytesIO()
            with zipfile.ZipFile(zip_buffer, "w") as zip_out:
                for file in os.listdir(word_dir):
                    full_path = os.path.join(word_dir, file)
                    arcname = os.path.join("Word Documents", file)
                    zip_out.write(full_path, arcname=arcname)

            st.success("✅ Word documents generated!")
            st.download_button(
                label="📦 Download All (Word Only – PDF not supported on Streamlit Cloud)",
                data=zip_buffer.getvalue(),
                file_name="word_documents.zip",
                mime="application/zip"
            )

    if template_mode == "Upload New Template":
        uploaded_template = st.file_uploader("Upload a .docx Template", type="docx")
        campaign_name = st.selectbox("🏷️ Select Campaign for This Template", CAMPAIGN_OPTIONS)
        doc_type = st.text_input("📄 Enter Document Type (e.g., HIPAA, Notice, Demand)")
        excel_file = st.file_uploader("Upload Excel Data (.xlsx)", type="xlsx", key="excel_upload_new")
        output_name_format = st.text_input("Enter filename format (e.g., HIPAA Notice ({{Client Name}}))")

        if uploaded_template and campaign_name and doc_type:
            if st.button("Save and Generate"):
                campaign_safe = campaign_name.replace(" ", "").replace("/", "-")
                doc_type_safe = doc_type.replace(" ", "")
                base_name = f"TEMPLATE_{doc_type_safe}_{campaign_safe}"
                version = 1
                while os.path.exists(os.path.join(TEMPLATE_FOLDER, f"{base_name}_v{version}.docx")):
                    version += 1
                final_filename = f"{base_name}_v{version}.docx"
                save_path = os.path.join(TEMPLATE_FOLDER, final_filename)

                with open(save_path, "wb") as f:
                    f.write(uploaded_template.read())

                st.success(f"✅ Saved as {final_filename}")
                if excel_file and output_name_format:
                    df = pd.read_excel(excel_file)
                    process_and_preview(save_path, df, output_name_format)

    elif template_mode == "Select a Saved Template":
        st.subheader("📂 Select a Saved Template")
        excluded_templates = {"foia_template.docx", "demand_template.docx"}
        available_templates = [
            f for f in os.listdir(TEMPLATE_FOLDER)
            if f.endswith(".docx") and f not in excluded_templates
        ]

        search_query = st.text_input("🔍 Search templates by keyword or campaign").lower()
        filtered_templates = [f for f in available_templates if search_query in f.lower()]

        if not filtered_templates:
            st.warning("⚠️ No matching templates found.")
            st.stop()

        template_choice = st.selectbox("Choose Template", filtered_templates)
        template_path = os.path.join(TEMPLATE_FOLDER, template_choice)
        excel_file = st.file_uploader("Upload Excel Data (.xlsx)", type="xlsx", key="excel_upload_saved")
        output_name_format = st.text_input("Enter filename format (e.g., HIPAA Notice ({{Client Name}}))")

        if st.button("Generate Documents"):
            if excel_file and output_name_format:
                df = pd.read_excel(excel_file)
                process_and_preview(template_path, df, output_name_format)

    elif template_mode == "Template Options":
        st.subheader("⚙️ Template Options")
        excluded_templates = {"foia_template.docx", "demand_template.docx"}
        available_templates = [
            f for f in os.listdir(TEMPLATE_FOLDER)
            if f.endswith(".docx") and f not in excluded_templates
        ]

        search_query = st.text_input("🔍 Search for template to manage").lower()
        filtered_templates = [f for f in available_templates if search_query in f.lower()]

        if filtered_templates:
            template_choice = st.selectbox("Choose Template to Rename/Delete", filtered_templates)
            template_path = os.path.join(TEMPLATE_FOLDER, template_choice)

            st.subheader("✏️ Rename Template")
            new_template_name = st.text_input("New name (no extension)", value=template_choice.replace(".docx", ""))
            if st.button("Rename Template"):
                new_path = os.path.join(TEMPLATE_FOLDER, new_template_name + ".docx")
                if os.path.exists(new_path):
                    st.warning("⚠️ A file with that name already exists.")
                else:
                    os.rename(template_path, new_path)
                    st.success(f"✅ Renamed to {new_template_name}.docx")
                    st.rerun()

            st.subheader("🗑️ Delete Template")
            confirm_delete = st.checkbox("Yes, delete this template permanently.")
            if st.button("Delete Template") and confirm_delete:
                os.remove(template_path)
                st.success(f"✅ Deleted '{template_choice}'")
                st.rerun()

        else:
            st.warning("⚠️ No templates found matching your search.")

# === Load Excel Data from Dropbox App Folder (Secure API Method) ===
if tool == "📊 Litigation Dashboard":
    from streamlit_autorefresh import st_autorefresh
    st.header("📊 Interactive Litigation Dashboard")
    st_autorefresh(interval=3600000, key="refresh_dashboard")

    try:
        # Authenticate with Dropbox using your secret token
        dbx = dropbox.Dropbox(st.secrets["dropbox_token"])
        file_path = "/Master Dashboard.xlsx"
        metadata, res = dbx.files_download(file_path)
        df = pd.read_excel(BytesIO(res.content), sheet_name="Master Dashboard")

        # === Sidebar Filters ===
        with st.sidebar:
            st.markdown("### 🔎 Filter by:")

            case_types = sorted(df["Case Type"].dropna().unique())
            class_codes = sorted(df["Class Code Title"].dropna().unique())
            referred_by = sorted(df["Referred By Name (Full - Last, First)"].dropna().unique())

            selected_case_types = st.multiselect("Case Type", case_types, default=case_types)
            selected_class_codes = st.multiselect("Class Code Title", class_codes, default=class_codes)
            selected_referrers = st.multiselect("Referred By", referred_by, default=referred_by)

        # === Apply Filters ===
        filtered_df = df[
            df["Case Type"].isin(selected_case_types) &
            df["Class Code Title"].isin(selected_class_codes) &
            df["Referred By Name (Full - Last, First)"].isin(selected_referrers)
        ]

        # === Display Results ===
        st.markdown(f"### 📁 Showing {len(filtered_df)} Case(s)")
        st.dataframe(filtered_df, use_container_width=True)

        # === Metrics Summary ===
        st.markdown("### 📊 Summary")
        col1, col2, col3 = st.columns(3)
        col1.metric("Cases Shown", len(filtered_df))
        col2.metric("Unique Referrers", filtered_df['Referred By Name (Full - Last, First)'].nunique())
        col3.metric("Case Types", filtered_df['Case Type'].nunique())

        # === Counts by Referrer and Class Code ===
        st.markdown("### 📊 Counts by Referrer")
        ref_counts = filtered_df['Referred By Name (Full - Last, First)'].value_counts().reset_index()
        ref_counts.columns = ["Referred By", "# of Cases"]
        st.dataframe(ref_counts, use_container_width=True)

        st.markdown("### 📊 Counts by Class Code")
        class_counts = filtered_df['Class Code Title'].value_counts().reset_index()
        class_counts.columns = ["Class Code Title", "# of Cases"]
        st.dataframe(class_counts, use_container_width=True)

    except Exception as e:
        st.error(f"❌ Could not load dashboard: {e}")
        st.stop()

# === Mediation Memo Generator (Simplified Input, Now with Multi-Depo Quote Support) ===
elif tool == "🧾 Mediation Memos":
    st.header("🧾 Generate Confidential Mediation Memo")

    if "depositions" not in st.session_state:
        st.session_state.depositions = []
    if "deposition_names" not in st.session_state:
        st.session_state.deposition_names = []
    if "quote_outputs" not in st.session_state:
        st.session_state.quote_outputs = {"Liability": [], "Damages": []}

    st.subheader("📘 Case Synopsis & Instructions")
    case_synopsis = st.text_area("💼 Brief Case Synopsis (optional)", height=100)
    quote_instructions = st.text_area(
        "📝 Instructions for AI (optional)",
        placeholder="E.g., Focus on quotes about emotional distress or negligent supervision...",
        height=100
    )

    st.subheader("📎 Add Deposition Excerpts One at a Time")

    new_depo_label = st.text_input("📏 Deposition Label (e.g., Efimov Deposition)")
    new_depo_text = st.text_area("✍️ Paste New Deposition Text", height=300)

    if st.button("➕ Add Deposition"):
        if new_depo_label.strip() and new_depo_text.strip():
            st.session_state.depositions.append(new_depo_text.strip())
            st.session_state.deposition_names.append(new_depo_label.strip())
            st.success(f"✅ '{new_depo_label.strip()}' added as Deposition #{len(st.session_state.depositions)}.")
        else:
            st.warning("Please enter both a label and deposition text.")


    if st.session_state.depositions:
        st.markdown("✅ **Depositions Loaded:**")
        for i, (depo, name) in enumerate(zip(st.session_state.depositions, st.session_state.deposition_names), 1):
            st.text_area(f"{name} (Deposition {i})", depo, height=150)


        if st.button("🧠 Extract Quotes from All Depositions"):
            from scripts.run_mediation import safe_generate, generate_with_openai
            st.session_state.quote_outputs = {"Liability": [], "Damages": []}

            for i, (depo_text, depo_name) in enumerate(zip(st.session_state.depositions, st.session_state.deposition_names), 1):
                with st.spinner(f"Analyzing Deposition #{i}..."):
                    prompt = f"""
You are a legal analyst reviewing deposition excerpts in a {case_synopsis.strip() or 'civil lawsuit'}.

Extract only **relevant Q&A quote pairs** that support **either LIABILITY or DAMAGES**.
Ignore all other content.

🧾 **Format for each Q&A**:
0012:24 Q: "What were you doing that day?"
0012:25 A: "I was monitoring intake procedures."

⚠️ **Rules**:
- Only include categories: **Liability** and **Damages**
- Return each quote under the appropriate section.
- Do not summarize. No headers or explanations.
- Use exact wording from the transcript. No edits.
- Show Q and A both, with full quote and line numbers.

{f"💡 Case Notes: {quote_instructions.strip()}" if quote_instructions.strip() else ""}

📄 Deposition:
{depo_text}
"""
                    try:
                        result = safe_generate(generate_with_openai, prompt, model="gpt-3.5-turbo")
                        if "**Damages**" in result:
                            liability_part, damages_part = result.split("**Damages**", 1)
                            if liability_part.strip():
                                labeled_liability = f"📑 {depo_name}\n{liability_part.strip()}"
                                st.session_state.quote_outputs["Liability"].append(
                                    f"📑 {st.session_state.deposition_names[i-1]}\n{result['liability_quotes']}"
                                )
                                st.session_state.quote_outputs["Damages"].append(
                                    f"📑 {st.session_state.deposition_names[i-1]}\n{result['damages_quotes']}"
                                )


                            if damages_part.strip():
                                labeled_damages = f"📑 {depo_name}\n{damages_part.strip()}"
                                st.session_state.quote_outputs["Damages"].append(labeled_damages)

                        else:
                            st.session_state.quote_outputs["Liability"].append(result.strip())
                    except Exception as e:
                        st.error(f"Error processing Deposition {i}: {e}")

        st.subheader("📂 Extracted Liability Quotes")
        st.text_area("Copy-ready Liability Quotes", "\n\n".join(st.session_state.quote_outputs["Liability"]), height=300)

        st.subheader("📂 Extracted Damages Quotes")
        st.text_area("Copy-ready Damages Quotes", "\n\n".join(st.session_state.quote_outputs["Damages"]), height=300)

    # === Memo Form (same as before) ===
    with st.form("simple_mediation_form"):
        court = st.text_input("Court")
        case_number = st.text_input("Case Number")

        plaintiffs = {}
        for i in range(1, 4):
            label = f"Plaintiff {i} Name" + (" (required)" if i == 1 else " (optional)")
            plaintiffs[f"plaintiff{i}"] = st.text_input(label)

        defendants = {}
        for i in range(1, 8):
            label = f"Defendant {i} Name" + (" (optional)" if i > 1 else "")
            defendants[f"defendant{i}"] = st.text_input(label)

        complaint_narrative = st.text_area("📔 Complaint Narrative", height=200)
        party_info = st.text_area("Party Information from Complaint", height=200)
        settlement_summary = st.text_area("💼 Settlement Demand Summary", height=200)
        medical_summary = st.text_area("🏥 Medical Summary", height=200)
        explicit_instructions = st.text_area("📝 Additional Instructions for Memo (optional)", height=100)

        deposition_liability = "\n\n".join(st.session_state.quote_outputs["Liability"])
        deposition_damages = "\n\n".join(st.session_state.quote_outputs["Damages"])

        submitted = st.form_submit_button("Generate Memo")

    if submitted:
        try:
            output_dir = "outputs/mediation_memos"
            os.makedirs(output_dir, exist_ok=True)

            data = {
                "court": court,
                "case_number": case_number,
                "complaint_narrative": complaint_narrative,
                "party_info": party_info,
                "settlement_summary": settlement_summary,
                "medical_summary": medical_summary,
                "deposition_liability": deposition_liability,
                "deposition_damages": deposition_damages,
                **plaintiffs,
                **defendants,
                "extracted_quotes": deposition_liability + "\n\n" + deposition_damages,
                "all_quotes_pool": deposition_liability + "\n\n" + deposition_damages,
            }

            template_path = "templates/mediation_template.docx"

            from scripts.run_mediation import safe_generate, generate_introduction, generate_plaintiff_statement, generate_defendant_statement, generate_demand_section, generate_facts_liability_section, generate_causation_injuries, generate_additional_harms, generate_future_medical, generate_conclusion_section, fill_mediation_template

            progress_text = st.empty()
            progress_bar = st.progress(0)

            steps = [
                ("Generating Introduction...", "introduction"),
                ("Generating Plaintiff Statement...", "plaintiff_statement")
            ]

            for i in range(1, 8):
                def_name = data.get(f"defendant{i}")
                if def_name:
                    steps.append((f"Generating Defendant {i} Statement...", f"defendant{i}_statement"))

            steps += [
                ("Generating Demand Section...", "demand"),
                ("Generating Facts / Liability Section...", "facts_liability"),
                ("Generating Causation & Injuries...", "causation_injuries"),
                ("Generating Additional Harms...", "additional_harms"),
                ("Generating Future Medical Costs...", "future_bills"),
                ("Generating Conclusion...", "conclusion")
            ]

            memo_data = {
                "Court": court,
                "Case Number": case_number,
                "Plaintiff1": plaintiffs["plaintiff1"],
            }

            for i in range(2, 4):
                memo_data[f"Plaintiff{i}"] = data.get(f"plaintiff{i}", "")
                memo_data[f"Plaintiff{i} Statement"] = ""

            for i in range(1, 8):
                memo_data[f"Defendant{i}"] = data.get(f"defendant{i}", "")
                memo_data[f"Defendant{i} Statement"] = ""

            total = len(steps)
            for idx, (text, key) in enumerate(steps):
                progress_text.text(text)

                if key == "introduction":
                    memo_data[key] = safe_generate(generate_introduction, data["complaint_narrative"], data["plaintiff1"])
                elif key == "plaintiff_statement":
                    memo_data["Plaintiff1 Statement"] = safe_generate(generate_plaintiff_statement, data["complaint_narrative"], data["plaintiff1"])
                elif key.startswith("defendant") and key.endswith("_statement"):
                    i = key.replace("defendant", "").replace("_statement", "")
                    memo_data[f"Defendant{i} Statement"] = safe_generate(generate_defendant_statement, data["complaint_narrative"], data[f"defendant{i}"])
                elif key == "demand":
                    memo_data[key] = safe_generate(generate_demand_section, data["settlement_summary"], data["plaintiff1"])
                elif key == "facts_liability":
                    memo_data[key] = safe_generate(generate_facts_liability_section, data["complaint_narrative"], data["deposition_liability"])
                elif key == "causation_injuries":
                    memo_data[key] = safe_generate(generate_causation_injuries, data["medical_summary"])
                elif key == "additional_harms":
                    memo_data[key] = safe_generate(generate_additional_harms, data["medical_summary"], data["deposition_damages"])
                elif key == "future_bills":
                    memo_data[key] = safe_generate(generate_future_medical, data["medical_summary"], data["deposition_damages"])
                elif key == "conclusion":
                    memo_data[key] = safe_generate(generate_conclusion_section, data["settlement_summary"])

                progress_bar.progress((idx + 1) / total)

            file_path = fill_mediation_template(data | memo_data, template_path, output_dir)

            with open(file_path, "rb") as f:
                st.success("✅ Mediation memo generated!")
                st.download_button("🗅️ Download Mediation Memo", f, file_name=os.path.basename(file_path))

            # Fallback plaintext output
            plaintext_output = generate_plaintext_memo(memo_data)
            st.download_button("📄 Download Plaintext Version", plaintext_output, file_name="Mediation_Memo_Plaintext.txt")
            st.text_area("📝 Memo Preview (Plaintext)", plaintext_output, height=700)


        except Exception as e:
            st.error(f"❌ Error: {e}")

if tool == "📖 Instructions & Support":
    st.header("📘 Instructions & Support")

    with st.expander("📄 Batch Doc Generator – How to Use", expanded=True):
        st.markdown("""
Use this tool to **automatically generate documents in bulk** by merging a Word template with an Excel sheet.

**Step-by-step:**
1. **Upload a Word Template or Select an Existing Template**
   - If uploading a new template, use placeholders inside of the document like `{{ClientName}}`, `{{Date}}`, etc.
   - These placeholders should mirror what is at the top of your excel columns
   - Save your template for reuse — it’ll appear in the dropdown.
   - If selecting an existing template, choose 'Select a Saved Template' and search for the existing template. 

2. **Upload an Excel File**
   - Must have one row per document.
   - Column names must match the placeholders in your Word template.

3. **Preview the Data**
   - View the first row to confirm the placeholder match.

4. **Set Output Filename Format**
   - Use any column name inside `{{ }}`.
   - Example: `{{ClientName}}_Notice` → `JohnDoe_Notice.docx`.

5. **Generate Documents**
   - Click “Generate Files.”
   - Download a ZIP file with all Word documents.
        """)

    with st.expander("📬 FOIA Requests – How to Use", expanded=False):
        st.markdown("""
Use this tool to generate **individual FOIA request letters** using form fields you fill in manually.

**Step-by-step:**
1. **Fill Out the Form**
   - Enter details like Client ID, Recipient info, Synopsis, Requested Records, and any Explicit Instructions (Optional but typically helpful to establish scope).
   - Enter the case type (Not Neos case type, get specific. Ex: Motorcycle Accident), the Facility or System (Ex: Municipal Police Department, DCFS, etc.), and the Recipient Role (Ex. Responding Officers). 
   - All inputs are required unless marked optional.

2. **Click 'Generate FOIA Letter'**
   - A personalized Word document will be created.

3. **Download**
   - You’ll see a preview and a download button for the generated letter.
        """)

    with st.expander("📂 Demand Letters – How to Use", expanded=False):
        st.markdown("""
Use this tool to generate **individual demand letters** using a manual entry form.

**Step-by-step:**
1. **Fill Out the Form**
   - Enter the client’s name, defendant, incident date, location, summary, and damages (damages should be a dollar figure. Ex. $100,000 One Hundred Thousand Dollars).

2. **Click 'Generate Demand Letter'**
   - The app will insert your responses into a Word template.

3. **Download**
   - A finished letter will be available for download after approximately a minute.
        """)

    with st.expander("🧾 Mediation Memos – How to Use", expanded=False):
        st.markdown("""
Use this tool to generate **confidential mediation memorandums** from structured prompts.

**Step-by-step:**
1. **Fill Out Each Section**
   - Court name, case number, and all memo sections (Intro through Conclusion).
   - AI will convert each input into a professional paragraph in your memo.

2. **Click 'Generate Mediation Memo'**
   - A final Word document will be created with your inputs formatted and polished.

3. **Download**
   - You'll get a download button to retrieve the complete mediation memorandum.
        """)

    with st.expander("🚧 Complaints – Coming Soon", expanded=False):
        st.markdown("""
This section will automate formal legal complaint drafting.

**Planned Features:**
- Upload structured abuse data (e.g., facility, dates, abuse summary).
- Auto-generate civil complaints using AI + custom templates.
- Optional footnotes, exhibits, and institutional history blocks.
        """)

    with st.expander("🚧 Subpoenas – Coming Soon", expanded=False):
        st.markdown("""
This tool will allow you to auto-generate subpoena forms for institutions, agencies, and record holders.

**Planned Features:**
- Upload a list of targets with addresses and requested materials.
- Pre-fill standard subpoena templates.
- Batch generate, preview, and download signed copies.
        """)

    st.subheader("🐞 Report a Bug")
    with st.form("report_form"):
        issue = st.text_area("Describe the issue:")
        submitted = st.form_submit_button("Submit")
        if submitted:
            with open("error_reports.txt", "a", encoding="utf-8") as f:
                f.write(issue + "\n---\n")
            st.success("✅ Issue submitted. Thank you!")

st.markdown("""
<hr style="margin-top: 2rem;">
<div style="text-align: center; font-size: 0.85rem; color: gray;">
&copy; 2025 Stinar Gould Grieco & Hensley. All rights reserved.
</div>
""", unsafe_allow_html=True)