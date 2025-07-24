import os
from datetime import datetime
from docx import Document
from docx.table import _Cell
from docx.text.paragraph import Paragraph
from openai import OpenAI
from openpyxl import load_workbook

# === OpenAI Setup ===
api_key = os.environ["OPENAI_API_KEY"]
client = OpenAI(api_key=api_key)

def generate_with_openai(prompt):
    response = client.chat.completions.create(
        model="gpt-4",
        messages=[
            {"role": "system", "content": "You are a professional legal writer."},
            {"role": "user", "content": prompt}
        ],
        temperature=0.5
    )
    return response.choices[0].message.content.strip()

# === Prompt Fragments ===
NO_HALLUCINATION_NOTE = """Do not fabricate or assume any facts..."""
LEGAL_FLUENCY_NOTE = """Use the tone and clarity of a senior litigator..."""
NO_PASSIVE_LANGUAGE_NOTE = """Every sentence must use active voice..."""
BAN_PHRASES_NOTE = """Ban weak phrasing like “may,” “might,” “ongoing discomfort,” etc..."""

FULL_SAFETY_PROMPT = "\n\n".join([
    NO_HALLUCINATION_NOTE,
    LEGAL_FLUENCY_NOTE,
    NO_PASSIVE_LANGUAGE_NOTE,
    BAN_PHRASES_NOTE
])

EXAMPLE_DEMAND = """
The snowplow operator owed a duty of care to nearby pedestrians and failed to exercise basic precautions while maneuvering in a populated area. This breach of duty was the direct and proximate cause of Jane’s orthopedic trauma and emotional harm. Accordingly, liability is clear.
"""

# === Section Generators ===
def generate_brief_synopsis(summary, full_name):
    prompt = f"""{FULL_SAFETY_PROMPT}

Write a one-sentence introduction for a legal demand letter. Refer to the client only as {full_name}. Begin with a sentence like: "Our office represents {full_name}, who suffered..." and describe the injuries factually and concisely. Do not repeat the client's name. Do not use more than one sentence.

Summary:
{summary}
"""
    response = generate_with_openai(prompt)
    return response.split(".")[0].strip() + "."

def generate_demand(summary, first_name):
    prompt = f"""{FULL_SAFETY_PROMPT}

Use this structure:
- Duty
- Standard of Care
- Breach
- Causation
- Harm

Style Guide Example (do not copy content):
{EXAMPLE_DEMAND}

Avoid re-describing injuries or facts. Reference “the injuries described above” or terms like “orthopedic trauma.” Frame causation, breach, and standards of care clearly and assertively.

Client: {first_name}
Summary:
{summary}
"""
    return generate_with_openai(prompt)

def generate_damages(damages_text, first_name):
    prompt = f"""{FULL_SAFETY_PROMPT}

Write a formal legal damages section for {first_name}. Do not restate facts or injury details. Focus on:
- medical disruption,
- treatment scope,
- financial strain,
- emotional hardship.

Use grouped providers or categories (e.g., ambulance, surgery, PT). Do not itemize every service. Avoid casual phrasing. Use phrases like:
- “Based on available billing records...”
- “Projecting from current provider estimates...”
- “Ongoing physical and emotional complications...”

Damages:
{damages_text}
"""
    return generate_with_openai(prompt)

def generate_settlement_demand(summary, damages_text, first_name):
    prompt = f"""{FULL_SAFETY_PROMPT}

Draft a closing paragraph to justify a settlement demand. Do not restate the incident. Emphasize:
- factual consistency,
- anticipated corroboration,
- litigation risk.

End with: “We invite resolution of this matter without the need for formal litigation. Should you fail to respond by [insert date], we are prepared to proceed accordingly.”

Client: {first_name}

Summary:
{summary}

Damages:
{damages_text}
"""
    return generate_with_openai(prompt)

# === Placeholder Replacement ===
def replace_placeholders(doc, replacements):
    def replace_in_paragraph(paragraph: Paragraph):
        if not paragraph.runs:
            return
        full_text = paragraph.text
        for key, val in replacements.items():
            full_text = full_text.replace(key, val)
        paragraph.clear()
        paragraph.add_run(full_text)

    def replace_in_cell(cell: _Cell):
        for paragraph in cell.paragraphs:
            replace_in_paragraph(paragraph)

    for paragraph in doc.paragraphs:
        replace_in_paragraph(paragraph)

    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                replace_in_cell(cell)

# === Word Document Filler ===
def fill_template(data, template_path, output_path):
    doc = Document(template_path)

    full_name = data.get("Client Name", "").strip()
    first_name = full_name.split()[0] if full_name else "Client"
    incident_date = data.get("Incident Date", "")
    if isinstance(incident_date, datetime):
        incident_date = incident_date.strftime("%B %d, %Y")
    elif not incident_date:
        incident_date = "[Date Not Provided]"

    summary = data.get("Summary", "") or "[No summary provided.]"
    damages = data.get("Damages", "") or "[No damages provided.]"

    replacements = {
        "{{Client Name}}": full_name,
        "{{IncidentDate}}": incident_date,
        "{{Brief Synopsis}}": generate_brief_synopsis(summary, full_name),
        "{{Demand}}": generate_demand(summary, first_name),
        "{{Damages}}": generate_damages(damages, first_name),
        "{{Settlement Demand}}": generate_settlement_demand(summary, damages, first_name)
    }

    replace_placeholders(doc, replacements)

    filename = f"Demand_{full_name.replace(' ', '_')}_{datetime.today().strftime('%Y-%m-%d')}.docx"
    final_path = os.path.join(output_path, filename)
    doc.save(final_path)
    print(f"Generated: {final_path}")
    return final_path

# === Batch Generation ===
def run(df):
    output_dir = "outputs/demands"
    os.makedirs(output_dir, exist_ok=True)
    template_path = "templates/demand_template.docx"
    output_paths = []

    for _, row in df.iterrows():
        data = {
            "Client Name": row.get("Client Name", ""),
            "Incident Date": row.get("Incident Date", ""),
            "Summary": row.get("Summary", ""),
            "Damages": row.get("Damages", "")
        }
        if data["Client Name"].strip():
            path = fill_template(data, template_path, output_dir)
            output_paths.append(path)

    return output_paths

# === Excel Integration ===
def generate_all_demands(template_path, excel_path, output_dir):
    wb = load_workbook(excel_path)
    sheet = wb.active
    headers = [str(cell.value).strip() for cell in sheet[1]]
    os.makedirs(output_dir, exist_ok=True)

    for row in sheet.iter_rows(min_row=2, values_only=True):
        data = dict(zip(headers, row))
        fill_template(data, template_path, output_dir)

# === Main Execution ===
if __name__ == "__main__":
    TEMPLATE_PATH = "templates/demand_template.docx"
    EXCEL_PATH = "data_demand_requests.xlsx"
    OUTPUT_DIR = "outputs/demands"
    generate_all_demands(TEMPLATE_PATH, EXCEL_PATH, OUTPUT_DIR)
