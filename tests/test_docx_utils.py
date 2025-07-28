from utils.docx_utils import replace_text_in_docx_all
from docx import Document
import os
from core.usage_tracker import check_quota, decrement_quota

def test_docx_placeholder_replacement(tmp_path):
    test_docx = tmp_path / "template.docx"
    test_output = tmp_path / "output.docx"

    doc = Document()
    doc.add_paragraph("Hello {{ClientName}}. Your case ID is {{CaseID}}.")
    doc.save(test_docx)

    check_quota("openai_tokens", amount=1)
    replace_text_in_docx_all(test_docx, {"ClientName": "Jane Roe", "CaseID": "12345"}, test_output)
    decrement_quota("openai_tokens", amount=1)

    result_doc = Document(test_output)
    text = "\n".join(p.text for p in result_doc.paragraphs)

    assert "Jane Roe" in text
    assert "12345" in text
